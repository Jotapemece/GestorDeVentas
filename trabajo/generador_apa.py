# -*- coding: utf-8 -*-
"""Generador de documento de caso clínico con normas APA 7ª edición.

Produce trabajo/Documento_APA.docx a partir de los datos de contenido_apa.py.
Estructura exacta según caso.pdf (Lúes Connatal).
"""
import os
from docx import Document
from docx.shared import Cm, Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import contenido_apa as c

# ----------------------------- Configuración APA -----------------------------
MARGEN = Cm(2.54)
SANGRIA = Cm(1.27)
INTERLINEADO = 1.5
FUENTE = 'Times New Roman'
TAM = Pt(12)

RUTA_SALIDA = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Documento_APA.docx')

# Entradas del índice: (titulo, sangria_cm, bookmark_id, negrita)
INDICE = [
    ('OBJETIVOS', 0, 'bm_objetivos', True),
    ('HISTORIA CLÍNICA DE ENFERMERÍA', 0, 'bm_hc', True),
    ('RESUMEN DEL CASO Y ENFERMEDAD ACTUAL', 0, 'bm_resumen', True),
    ('EXÁMENES DE LABORATORIO Y PARACLÍNICOS', 0, 'bm_lab', True),
    ('DIAGNÓSTICO MÉDICO', 0, 'bm_dx', True),
    ('ANTECEDENTES PERSONALES PATOLÓGICOS', 0, 'bm_ante', True),
    ('EXAMEN FÍSICO', 0, 'bm_examen', True),
    ('PATRONES FUNCIONALES DE MARJORY GORDON', 0, 'bm_patrones', True),
    ('PLAN DE CUIDADOS E INTERVENCIONES DE ENFERMERÍA', 0, 'bm_plan', True),
    ('CONCLUSIONES Y RECOMENDACIONES', 0, 'bm_conclusion', True),
]


# ----------------------------- Helpers base -----------------------------
def _run(p, texto, negrita=False, cursiva=False, tam=TAM):
    r = p.add_run(texto)
    r.font.name = FUENTE
    r.font.size = tam
    r.font.bold = negrita
    r.font.italic = cursiva
    return r


def _activar_actualizacion_campos(doc):
    settings = doc.settings.element
    upd = settings.find(qn('w:updateFields'))
    if upd is None:
        upd = OxmlElement('w:updateFields')
        upd.set(qn('w:val'), 'true')
        settings.insert(0, upd)


def setup_apa(doc):
    sec = doc.sections[0]
    sec.top_margin = MARGEN
    sec.bottom_margin = MARGEN
    sec.left_margin = MARGEN
    sec.right_margin = MARGEN

    style = doc.styles['Normal']
    style.font.name = FUENTE
    style.font.size = TAM
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), FUENTE)
    rfonts.set(qn('w:hAnsi'), FUENTE)
    rfonts.set(qn('w:eastAsia'), FUENTE)
    pf = style.paragraph_format
    pf.line_spacing = INTERLINEADO
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


# ----------------------------- Párrafos y títulos -----------------------------
def parrafo_normado(doc, texto='', sangria=True, negrita=False, cursiva=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = INTERLINEADO
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = SANGRIA if sangria else Cm(0)
    if texto:
        _run(p, texto, negrita, cursiva)
    return p


def titulo_centrado(doc, texto, bookmark_id=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = INTERLINEADO
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    pf.first_line_indent = Cm(0)
    _run(p, texto.upper())
    if bookmark_id:
        _agregar_bookmark(p, bookmark_id)
    return p


def subtitulo_izq(doc, texto, bookmark_id=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = INTERLINEADO
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0)
    t = texto[0].upper() + (texto[1:] if texto else '') if texto else texto
    _run(p, t, negrita=True)
    if bookmark_id:
        _agregar_bookmark(p, bookmark_id)
    return p


def _render_parrafos(doc, parrafos, sangria=True):
    if not parrafos:
        parrafo_normado(doc, '', sangria=False)
    else:
        for txt in parrafos:
            parrafo_normado(doc, txt, sangria=sangria)


# ----------------------------- Bookmarks / PAGEREF -----------------------------
_bid = [0]


def _agregar_bookmark(parrafo, name):
    _bid[0] += 1
    pid = str(_bid[0])
    p_elem = parrafo._element
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), pid)
    start.set(qn('w:name'), name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), pid)
    end.set(qn('w:name'), name)
    p_elem.insert(0, start)
    p_elem.append(end)


def _agregar_campo_pageref(parrafo, bookmark_id, negrita=False, valor_calculado='?'):
    run = parrafo.add_run()
    run.font.name = FUENTE
    run.font.size = TAM
    run.font.bold = negrita
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' PAGEREF {bookmark_id} \\h '
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = valor_calculado
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_sep)
    run._element.append(t)
    run._element.append(fld_end)


def agregar_parada_tabulacion_puntos(parrafo, posicion_emu):
    parrafo.paragraph_format.tab_stops.add_tab_stop(
        Emu(posicion_emu),
        alignment=WD_TAB_ALIGNMENT.RIGHT,
        leader=WD_TAB_LEADER.DOTS,
    )


def agregar_fila_indice(doc, titulo, sangria_cm=0, bookmark_id=None, negrita=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = INTERLINEADO
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(sangria_cm)
    section = doc.sections[-1]
    ancho_util = section.page_width - section.left_margin - section.right_margin
    agregar_parada_tabulacion_puntos(p, ancho_util)
    run_desc = p.add_run(f'{titulo}\t')
    run_desc.font.name = FUENTE
    run_desc.font.size = TAM
    run_desc.font.bold = negrita
    if bookmark_id:
        _agregar_campo_pageref(p, bookmark_id, negrita=negrita, valor_calculado='?')
    else:
        run_pag = p.add_run('?')
        run_pag.font.name = FUENTE
        run_pag.font.size = TAM
        run_pag.font.bold = negrita
    return p


def agregar_numeracion_superior_der(doc, section=None):
    section = section or doc.sections[0]
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = p.paragraph_format
    pf.line_spacing = INTERLINEADO
    run = p.add_run()
    run.font.name = FUENTE
    run.font.size = TAM
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = '1'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_sep)
    run._element.append(t)
    run._element.append(fld_end)


# ----------------------------- Portada -----------------------------
def agregar_pie_portada(doc, texto):
    sec = doc.sections[0]
    footer = sec.first_page_footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(texto)
    r.font.name = FUENTE
    r.font.size = TAM


def _celda_sin_bordes(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        val = {'val': 'none'}
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), val.get('val', 'single'))
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), '000000')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def construir_portada(doc, datos):
    for linea in datos.get('institucion', []):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = INTERLINEADO
        _run(p, linea)

    for _ in range(4):
        parrafo_normado(doc, '', sangria=False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = INTERLINEADO
    r = p.add_run(datos.get('titulo', ''))
    r.font.name = FUENTE
    r.font.size = TAM
    r.font.bold = True

    for _ in range(8):
        parrafo_normado(doc, '', sangria=False)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = True
    left = tbl.cell(0, 0)
    right = tbl.cell(0, 1)
    _celda_sin_bordes(left)
    _celda_sin_bordes(right)

    pl = left.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = pl.add_run('TUTOR:')
    r1.font.name = FUENTE
    r1.font.size = TAM
    r1.font.bold = True
    pl.add_run().add_break()
    r2 = pl.add_run(datos.get('tutor', ''))
    r2.font.name = FUENTE
    r2.font.size = TAM

    pr = right.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = pr.add_run('INTEGRANTES:')
    r3.font.name = FUENTE
    r3.font.size = TAM
    r3.font.bold = True
    for autor in datos.get('autores', []):
        pr.add_run().add_break()
        r4 = pr.add_run(autor)
        r4.font.name = FUENTE
        r4.font.size = TAM


# ----------------------------- Ensamblaje -----------------------------
def main():
    doc = Document()
    setup_apa(doc)
    _activar_actualizacion_campos(doc)
    agregar_numeracion_superior_der(doc)

    # ---- Portada ----
    doc.sections[0].different_first_page_header_footer = True
    construir_portada(doc, c.PORTADA)
    if c.PORTADA.get('fecha'):
        agregar_pie_portada(doc, c.PORTADA['fecha'])
    doc.add_page_break()

    # ---- Preliminares ----
    if c.AGRADECIMIENTOS:
        titulo_centrado(doc, 'AGRADECIMIENTOS', 'bm_agrade')
        _render_parrafos(doc, c.AGRADECIMIENTOS, sangria=False)
        doc.add_page_break()
    if c.DEDICATORIA:
        titulo_centrado(doc, 'DEDICATORIA', 'bm_ded')
        _render_parrafos(doc, c.DEDICATORIA, sangria=False)
        doc.add_page_break()

    # ---- Tabla de Contenido ----
    titulo_centrado(doc, 'TABLA DE CONTENIDO', 'bm_indice')
    for titulo, sangria_cm, bm, neg in INDICE:
        agregar_fila_indice(doc, titulo, sangria_cm, bm, neg)
    doc.add_page_break()

    # ---- Introducción ----
    if c.INTRODUCCION:
        titulo_centrado(doc, 'INTRODUCCIÓN', 'bm_intro')
        _render_parrafos(doc, c.INTRODUCCION)
        doc.add_page_break()

    # ---- OBJETIVOS ----
    titulo_centrado(doc, 'OBJETIVOS', 'bm_objetivos')
    subtitulo_izq(doc, 'Objetivo General')
    if c.OBJETIVO_GENERAL:
        parrafo_normado(doc, c.OBJETIVO_GENERAL)
    subtitulo_izq(doc, 'Objetivos Específicos')
    for i, obj in enumerate(c.OBJETIVOS_ESPECIFICOS, 1):
        p = parrafo_normado(doc, sangria=False)
        _run(p, f'{i}. {obj}')
    doc.add_page_break()

    # ---- HISTORIA CLÍNICA DE ENFERMERÍA ----
    titulo_centrado(doc, 'HISTORIA CLÍNICA DE ENFERMERÍA', 'bm_hc')
    _render_parrafos(doc, c.HISTORIA_CLINICA, sangria=False)
    doc.add_page_break()

    # ---- RESUMEN DEL CASO Y ENFERMEDAD ACTUAL ----
    titulo_centrado(doc, 'RESUMEN DEL CASO Y ENFERMEDAD ACTUAL', 'bm_resumen')
    _render_parrafos(doc, c.RESUMEN_CASO)
    doc.add_page_break()

    # ---- EXÁMENES DE LABORATORIO Y PARACLÍNICOS ----
    titulo_centrado(doc, 'EXÁMENES DE LABORATORIO Y PARACLÍNICOS', 'bm_lab')
    _render_parrafos(doc, c.EXAMENES_LABORATORIO, sangria=False)
    doc.add_page_break()

    # ---- DIAGNÓSTICO MÉDICO (DEFINICIÓN Y AFECTACIÓN CLÍNICA) ----
    titulo_centrado(doc, 'DIAGNÓSTICO MÉDICO (DEFINICIÓN Y AFECTACIÓN CLÍNICA)', 'bm_dx')
    _render_parrafos(doc, c.DX_MEDICO)
    doc.add_page_break()

    # ---- ANTECEDENTES PERSONALES PATOLÓGICOS (MATERNOS / PERINATALES) ----
    titulo_centrado(doc, 'ANTECEDENTES PERSONALES PATOLÓGICOS (MATERNOS / PERINATALES)', 'bm_ante')
    _render_parrafos(doc, c.ANTECEDENTES)
    doc.add_page_break()

    # ---- EXAMEN FÍSICO (CEFALOCAUDAL) ----
    titulo_centrado(doc, 'EXAMEN FÍSICO (CEFALOCAUDAL)', 'bm_examen')
    _render_parrafos(doc, c.EXAMEN_FISICO, sangria=False)
    doc.add_page_break()

    # ---- PATRONES FUNCIONALES DE MARJORY GORDON (ALTERADOS) ----
    titulo_centrado(doc, 'PATRONES FUNCIONALES DE MARJORY GORDON (ALTERADOS)', 'bm_patrones')
    _render_parrafos(doc, c.PATRONES_GORDON)
    doc.add_page_break()

    # ---- PLAN DE CUIDADOS E INTERVENCIONES DE ENFERMERÍA (5 DIAGNÓSTICOS NANDA) ----
    titulo_centrado(doc, 'PLAN DE CUIDADOS E INTERVENCIONES DE ENFERMERÍA (5 DIAGNÓSTICOS NANDA)', 'bm_plan')
    _render_parrafos(doc, c.PLAN_CUIDADOS)
    doc.add_page_break()

    # ---- CONCLUSIONES Y RECOMENDACIONES (PERSONALES) ----
    titulo_centrado(doc, 'CONCLUSIONES Y RECOMENDACIONES (PERSONALES)', 'bm_conclusion')
    if c.CONCLUSIONES:
        subtitulo_izq(doc, 'Conclusiones')
        _render_parrafos(doc, c.CONCLUSIONES)
    if c.RECOMENDACIONES:
        subtitulo_izq(doc, 'Recomendaciones')
        _render_parrafos(doc, c.RECOMENDACIONES)
    if c.QUE_APRENDI:
        doc.add_page_break()
        titulo_centrado(doc, '¿QUÉ APRENDÍ EN LAS PASANTÍAS?', 'bm_qaprendi')
        _render_parrafos(doc, c.QUE_APRENDI)

    doc.save(RUTA_SALIDA)
    print('OK ->', RUTA_SALIDA)


if __name__ == '__main__':
    main()
