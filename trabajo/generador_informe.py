import os
import math
import subprocess
import platform
from docx import Document
from docx.shared import Cm, Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION_START, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Importamos el contenido
import contenido as c

# ================================================================
#  CONSTANTES DE FORMATO (NORMATIVA IUTECP)
# ================================================================

# Fuente y tamaño (Art. 5)
FUENTE = 'Times New Roman'
TAMANO_BASE = 12           # Pt
TAMANO_TABLA = 10          # Pt (cuerpo de tablas)
TAMANO_TABLA_CHICO = 9     # Pt (encabezados compactos Gantt)

# Página (Art. 6)
PAG_ANCHO = Cm(21.59)     # Carta
PAG_ALTO  = Cm(27.94)     # Carta
MARGEN_IZQ = Cm(4)        # Encuadernación
MARGEN_DER = Cm(3)
MARGEN_SUP = Cm(3)
MARGEN_INF = Cm(3)
MARGEN_SUP_CAP = Cm(5)    # Primera página de capítulo/parte

# Tipografía
INTERLINEADO = 1.5
SANGRIA_LINEA = Cm(1.25)          # Sangría de primera línea (Art. 7)
SANGRIA_CITA = Cm(1.25)           # Citas largas (Art. 22)
SANGRIA_NIVEL45 = Cm(1.27)        # Niveles 4 y 5 (Art. 9)
SANGRIA_REF = Cm(0.75)            # Sangría francesa referencias (Art. 25)

# Espaciados (Art. 8)
ESP_DOBLE = Pt(24)
ESP_SENCILLO = Pt(12)
ESP_TRIPLE = Pt(36)
ESP_CITA_ANTES = Pt(36)           # 3 espacios antes de cita larga
ESP_CITA_DESPUES = Pt(12)

# Colores de tablas
COLOR_ENCABEZADO = '4472C4'       # Azul corporativo
COLOR_FILA_PAR = 'D9E1F2'         # Gris azulado
COLOR_GANTT_VERDE = '70AD47'      # Barras activas Gantt (trabajo operativo)
COLOR_GANTT_INVESTIGACION = 'ED7D31'  # Barras activas Gantt (proyecto/investigación)
COLOR_GANTT_MIXTA = '7030A0'       # Barras activas Gantt (operativo + proyecto)
COLOR_TEXTO_CLARO = 'FFFFFF'
COLOR_TEXTO_OSCURO = '000000'

# Producto: nombres de archivo de salida
DOCX_SALIDA = "Informe_Pasantia_IUTECP.docx"
PDF_SALIDA  = "Informe_Pasantia_IUTECP.pdf"

# Espaciado de la portadilla de ANEXOS (Art. 26)
ESP_PORTADILLA_ANEXOS = Pt(237)  # 180pt + 57pt (compensa cambio de 5cm a 3cm)
TAMANO_PORTADILLA_ANEXOS = 16  # Pt, un poco más grande para portadilla

# Títulos por defecto de Cuadros y Gráficos (configurables desde contenido.py)
CUADRO_PLANIFICACION_TITULO_DEF = "Cuadro 1. Planificación integral de objetivos específicos."
CUADRO_CRONOGRAMA_TITULO_DEF = "Cuadro 2. Cronograma de actividades administrativas."

# Parágrafo posterior a la cita (configurable desde contenido.py)
POST_CITA_TEXTO_DEF = (
    "De acuerdo a la cita previa, se comprende la relevancia del control sistemático "
    "y la inmutabilidad de los registros en los departamentos estratégicos de la empresa."
)

# ================================================================
#  FUNCIONES AUXILIARES DE FORMATO (NORMATIVA IUTECP 2025)
# ================================================================

def setup_iutecp_document():
    """Configura el documento base: Tamaño Carta, Márgenes Art. 6, Fuente Art. 5"""
    doc = Document()

    section = doc.sections[0]
    section.page_width = PAG_ANCHO
    section.page_height = PAG_ALTO
    section.top_margin = MARGEN_SUP
    section.bottom_margin = MARGEN_INF
    section.left_margin = MARGEN_IZQ      # Encuadernación
    section.right_margin = MARGEN_DER

    style = doc.styles['Normal']
    style.font.name = FUENTE
    style.font.size = Pt(TAMANO_BASE)
    style.paragraph_format.line_spacing = INTERLINEADO
    style.paragraph_format.space_after = Pt(0)

    # Pedir a Word/LibreOffice que recalculen campos (PAGEREF del índice, PAGE) al abrir
    _activar_actualizacion_campos(doc)

    return doc

def _activar_actualizacion_campos(doc):
    """Inserta <w:updateFields w:val="true"/> para que los visores recalculen
    los campos PAGE/PAGEREF (índice, listas) al abrir el documento."""
    settings = doc.settings.element
    upd = settings.find(qn('w:updateFields'))
    if upd is None:
        upd = OxmlElement('w:updateFields')
        upd.set(qn('w:val'), 'true')
        settings.insert(0, upd)

def agregar_parrafo_normado(doc, texto, cursiva=False, sangria=True):
    """Párrafo justificado, 1.5 interlineado, sangría 1.25cm (Art. 7, 8)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_format = p.paragraph_format
    p_format.line_spacing = INTERLINEADO
    p_format.first_line_indent = SANGRIA_LINEA if sangria else Cm(0)
    p_format.space_after = Pt(0)
    p_format.space_before = Pt(0)
    p_format.keep_together = True
    p_format.widow_control = True

    run = p.add_run(texto)
    run.font.name = FUENTE
    run.font.size = Pt(TAMANO_BASE)
    run.font.italic = cursiva
    return p

def agregar_item_lista(doc, numero, texto, negrita_inicio=""):
    """Agrega un elemento enumerado con sangría de primera línea (Art. 10)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_format = p.paragraph_format
    p_format.line_spacing = INTERLINEADO
    p_format.first_line_indent = SANGRIA_LINEA
    p_format.keep_together = True
    p_format.widow_control = True

    run_num = p.add_run(f"{numero}. ")
    run_num.font.name = FUENTE
    run_num.font.size = Pt(TAMANO_BASE)

    if negrita_inicio:
        run_bold = p.add_run(f"{negrita_inicio}: ")
        run_bold.font.name = FUENTE
        run_bold.font.size = Pt(TAMANO_BASE)
        run_bold.font.bold = True

    run_text = p.add_run(texto)
    run_text.font.name = FUENTE
    run_text.font.size = Pt(TAMANO_BASE)
    return p

def agregar_viñeta(doc, texto):
    """Elemento con viñeta (-) y sangría de primera línea para listados (Art. 10)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_format = p.paragraph_format
    p_format.line_spacing = INTERLINEADO
    p_format.first_line_indent = SANGRIA_LINEA
    p_format.keep_together = True
    p_format.widow_control = True

    run_vin = p.add_run("• ")
    run_vin.font.name = FUENTE
    run_vin.font.size = Pt(TAMANO_BASE)

    run_text = p.add_run(texto)
    run_text.font.name = FUENTE
    run_text.font.size = Pt(TAMANO_BASE)
    return p

def agregar_titulo_nivel2(doc, texto, bookmark_id=None):
    """Nivel 2: Alineado a la izquierda, Negrita (Art. 9 - IUTECP)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = ESP_DOBLE
    p.paragraph_format.space_after = ESP_DOBLE
    p.paragraph_format.line_spacing = INTERLINEADO
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True

    run = p.add_run(texto)
    run.font.name = FUENTE
    run.font.size = Pt(TAMANO_BASE)
    run.font.bold = True
    if bookmark_id:
        _agregar_bookmark(p, bookmark_id)
    return p

def agregar_parrafo_seccion(doc, etiqueta, texto):
    """Párrafo justificado con etiqueta en negrita al inicio y cuerpo normal,
    formato idéntico al PDF de referencia del Capítulo IV.
    (Ej.: "Operativa: texto..." con "Operativa" en negrita)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_format = p.paragraph_format
    p_format.line_spacing = INTERLINEADO
    p_format.first_line_indent = SANGRIA_LINEA
    p_format.space_after = Pt(0)
    p_format.space_before = Pt(0)
    p_format.keep_together = True
    p_format.widow_control = True

    if etiqueta:
        run_et = p.add_run(f"{etiqueta}: ")
        run_et.font.name = FUENTE
        run_et.font.size = Pt(TAMANO_BASE)
        run_et.font.bold = True

    run_text = p.add_run(texto)
    run_text.font.name = FUENTE
    run_text.font.size = Pt(TAMANO_BASE)
    return p

def agregar_titulo_nivel3(doc, texto, bookmark_id=None):
    """Nivel 3: Alineado a la izquierda, Negrita y Cursiva (Art. 9 - IUTECP)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = ESP_SENCILLO
    p.paragraph_format.space_after = ESP_SENCILLO
    p.paragraph_format.line_spacing = INTERLINEADO
    p.paragraph_format.first_line_indent = SANGRIA_LINEA
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True

    run = p.add_run(texto)
    run.font.name = FUENTE
    run.font.size = Pt(TAMANO_BASE)
    run.font.bold = True
    run.font.italic = True
    if bookmark_id:
        _agregar_bookmark(p, bookmark_id)
    return p

def agregar_titulo_nivel4(doc, texto, texto_parrafo=''):
    """
    Nivel 4: Izquierda · Negrita · Sangría ½ pulgada (1.27 cm) · Punto final
    El encabezado termina en punto y el texto del párrafo continúa en la misma línea.
    (Art. 9 - IUTECP)
    Parámetros:
        texto         : El título del nivel 4 (sin punto, se agrega automáticamente).
        texto_parrafo : Texto que continúa en la misma línea después del título.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = ESP_SENCILLO
    p.paragraph_format.space_after = ESP_SENCILLO
    p.paragraph_format.line_spacing = INTERLINEADO
    p.paragraph_format.left_indent = SANGRIA_NIVEL45
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True

    # Encabezado en negrita con punto final
    run_header = p.add_run(texto.rstrip('.') + '. ')
    run_header.font.name = FUENTE
    run_header.font.size = Pt(TAMANO_BASE)
    run_header.font.bold = True

    # Texto del párrafo en la misma línea (formato normal)
    if texto_parrafo:
        run_body = p.add_run(texto_parrafo)
        run_body.font.name = FUENTE
        run_body.font.size = Pt(TAMANO_BASE)
    return p

def agregar_titulo_nivel5(doc, texto, texto_parrafo=''):
    """
    Nivel 5: Izquierda · Negrita · Cursiva · Sangría ½ pulgada (1.27 cm) · Punto final
    El encabezado termina en punto y el texto del párrafo continúa en la misma línea.
    (Art. 9 - IUTECP)
    Parámetros:
        texto         : El título del nivel 5 (sin punto, se agrega automáticamente).
        texto_parrafo : Texto que continúa en la misma línea después del título.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True

    # Encabezado en negrita + cursiva con punto final
    run_header = p.add_run(texto.rstrip('.') + '. ')
    run_header.font.name = 'Times New Roman'
    run_header.font.size = Pt(12)
    run_header.font.bold = True
    run_header.font.italic = True

    # Texto del párrafo en la misma línea (formato normal)
    if texto_parrafo:
        run_body = p.add_run(texto_parrafo)
        run_body.font.name = 'Times New Roman'
        run_body.font.size = Pt(12)
    return p

def _config_seccion(sec, sup=None, primera_sin_numero=True):
    """Aplica los márgenes IUTECP a una sección (Art. 6). `sup` override del margen superior.
    `primera_sin_numero`: si True, activa títulos de página (titlePg) para que la primera
    página de la sección no muestre el número (solo en arranques de capítulo/preliminar)."""
    sec.top_margin = sup if sup is not None else MARGEN_SUP
    sec.bottom_margin = MARGEN_INF
    sec.left_margin = MARGEN_IZQ
    sec.right_margin = MARGEN_DER
    sec.different_first_page_header_footer = primera_sin_numero

def iniciar_capitulo(doc, numero_romano, titulo, bookmark_id=None):
    """
    Crea un nuevo capítulo con espacio extra de 2cm en el título (simula Art. 6),
    centrado, negrita, mayúsculas (Art. 9).
    """
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _config_seccion(sec)

    # Título CAPÍTULO X — con espacio extra para simular margen superior de 5cm
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(57)
    p1.paragraph_format.space_after = ESP_DOBLE
    r1 = p1.add_run(f"CAPÍTULO {numero_romano}")
    r1.font.name = FUENTE
    r1.font.size = Pt(TAMANO_BASE)
    r1.font.bold = True

    # Título del Capítulo
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = ESP_DOBLE
    r2 = p2.add_run(titulo.upper())
    r2.font.name = FUENTE
    r2.font.size = Pt(TAMANO_BASE)
    r2.font.bold = True
    if bookmark_id:
        _agregar_bookmark(p2, bookmark_id)

    sec2 = doc.add_section(WD_SECTION_START.CONTINUOUS)
    _config_seccion(sec2, primera_sin_numero=False)

def iniciar_seccion_preliminar(doc, titulo, bookmark_id=None, ocultar_primera=False):
    """Para secciones preliminares que inician en página nueva (Art. 9, 12).

    `ocultar_primera=True` deja sin número la primera página de la sección
    (usado solo en la INTRODUCCIÓN y en los capítulos). Las demás preliminares
    muestran su número romano en todas sus páginas.
    """
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _config_seccion(sec, primera_sin_numero=ocultar_primera)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(57)
    p.paragraph_format.space_after = ESP_DOBLE
    r = p.add_run(titulo.upper())
    r.font.name = FUENTE
    r.font.size = Pt(TAMANO_BASE)
    r.font.bold = True
    if bookmark_id:
        _agregar_bookmark(p, bookmark_id)

    sec2 = doc.add_section(WD_SECTION_START.CONTINUOUS)
    _config_seccion(sec2, primera_sin_numero=False)

def iniciar_seccion_resumen(doc, contenido, bookmark_id=None):
    """Construye la página de resumen con membrete, autor y fecha."""
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _config_seccion(sec, primera_sin_numero=False)

    for linea in contenido.MEMBRETE:
        p_membrete = doc.add_paragraph()
        p_membrete.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_membrete.paragraph_format.line_spacing = 1.0
        p_membrete.paragraph_format.space_before = Pt(0)
        p_membrete.paragraph_format.space_after = Pt(0)
        run_membrete = p_membrete.add_run(linea)
        run_membrete.font.name = FUENTE
        run_membrete.font.size = Pt(TAMANO_BASE)
        run_membrete.font.bold = True

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.line_spacing = 1.15
    p_titulo.paragraph_format.space_before = Pt(6)
    p_titulo.paragraph_format.space_after = Pt(6)
    run_titulo = p_titulo.add_run(contenido.TITULO_PROYECTO.upper())
    run_titulo.font.name = FUENTE
    run_titulo.font.size = Pt(11)
    run_titulo.font.bold = True

    fecha = getattr(contenido, 'FECHA_LUGAR', '')
    fecha_resumen = fecha.split(',', 1)[1].strip() if ',' in fecha else fecha
    p_autor = doc.add_paragraph()
    p_autor.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_autor.paragraph_format.left_indent = Cm(0)
    p_autor.paragraph_format.line_spacing = 1.0
    p_autor.paragraph_format.space_after = Pt(6)
    for indice, texto in enumerate((
        "Autor:",
        contenido.NOMBRE_PASANTE,
        f"C.I.: {contenido.CI_PASANTE}",
        fecha_resumen,
    )):
        run_autor = p_autor.add_run(texto)
        run_autor.font.name = FUENTE
        run_autor.font.size = Pt(TAMANO_BASE)
        if indice == 0:
            run_autor.font.bold = True
        if indice < 3:
            run_autor.add_break()

    p_resumen = doc.add_paragraph()
    p_resumen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_resumen.paragraph_format.space_before = Pt(6)
    p_resumen.paragraph_format.space_after = Pt(6)
    run_resumen = p_resumen.add_run("RESUMEN")
    run_resumen.font.name = FUENTE
    run_resumen.font.size = Pt(TAMANO_BASE)
    run_resumen.font.bold = True
    if bookmark_id:
        _agregar_bookmark(p_resumen, bookmark_id)

    sec2 = doc.add_section(WD_SECTION_START.CONTINUOUS)
    _config_seccion(sec2)

def agregar_cita_larga(doc, texto, cita):
    """
    Citas textuales de más de 40 palabras:
    Párrafo separado, sangría 1.25cm (5 espacios) a ambos lados, interlineado sencillo,
    sin comillas, distancia de 3 espacios de separación del párrafo anterior (Art. 7, 22).
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = SANGRIA_CITA
    p.paragraph_format.right_indent = SANGRIA_CITA
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = ESP_CITA_ANTES
    p.paragraph_format.space_after = ESP_CITA_DESPUES

    run_t = p.add_run(texto)
    run_t.font.name = FUENTE
    run_t.font.size = Pt(TAMANO_BASE)

    run_c = p.add_run(f" {cita}")
    run_c.font.name = FUENTE
    run_c.font.size = Pt(TAMANO_BASE)

def agregar_referencia(doc, texto):
    """
    Entrada en referencias bibliográficas (Art. 7, 25 IUTECP):
    - Interlineado sencillo dentro de cada entrada.
    - Sangría francesa de 3 espacios (~0.75cm) hacia la derecha.
    - Entre una referencia y otra: dos (2) espacios sencillos de separación.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = SANGRIA_REF
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.line_spacing = 1.0
    # 2 espacios sencillos de separación entre referencias (Art. 25)
    p.paragraph_format.space_after = ESP_DOBLE
    p.paragraph_format.space_before = Pt(0)

    run = p.add_run(texto)
    run.font.name = FUENTE
    run.font.size = Pt(TAMANO_BASE)

def set_cell_border(cell, **kwargs):
    """Permite definir bordes de celda personalizados mediante XML"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        val = kwargs.get(side, {'val': 'single', 'sz': '4', 'color': '000000'})
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   val.get('val', 'single'))
        el.set(qn('w:sz'),    val.get('sz', '4'))
        el.set(qn('w:color'), val.get('color', '000000'))
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_cell_shading(cell, fill='4472C4'):
    """Establece el color de fondo de una celda mediante XML"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    tcPr.append(shd)

def set_table_fixed_layout(tabla):
    """Establece layout de tabla fijo en XML"""
    tbl = tabla._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(existing)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

def set_cell_width(cell, width_cm):
    """Establece el ancho de una celda en XML (en dxa/twips)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcW')):
        tcPr.remove(existing)
    tcW = OxmlElement('w:tcW')
    twips = int(width_cm * 567)
    tcW.set(qn('w:w'), str(twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def set_table_grid_widths_xml(tabla, lista_anchos_cm):
    """Aplica el gridCol XML para garantizar renderizado exacto en LibreOffice/Word"""
    tbl = tabla._tbl
    tblGrid = tbl.tblGrid
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.append(tblGrid)
    else:
        for existing in tblGrid.findall(qn('w:gridCol')):
            tblGrid.remove(existing)

    for ancho_cm in lista_anchos_cm:
        gridCol = OxmlElement('w:gridCol')
        twips = int(ancho_cm * 567)
        gridCol.set(qn('w:w'), str(twips))
        tblGrid.append(gridCol)

def aplicar_formato_tabla_xml(tabla, lista_anchos_cm):
    """Aplica el ancho fijo a la tabla y a cada celda de forma inmutable"""
    set_table_fixed_layout(tabla)
    set_table_grid_widths_xml(tabla, lista_anchos_cm)
    for fila in tabla.rows:
        for col_idx, ancho in enumerate(lista_anchos_cm):
            set_cell_width(fila.cells[col_idx], ancho)

def _celda(cell, texto, negrita=False, centrado=False, tamaño=Pt(TAMANO_TABLA), color_texto=COLOR_TEXTO_OSCURO):
    """Inserta texto formateado en una celda de tabla limpiando párrafos vacíos.

    Si texto es una lista de cadenas (o una sola cadena con '\\n'), cada línea va en un
    párrafo aparte, permitiendo bullets multinivel en la misma celda.
    """
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)

    if isinstance(texto, (list, tuple)):
        lineas = [str(x) for x in texto]
    else:
        lineas = str(texto).split("\n")

    primer = True
    for linea in lineas:
        if not primer and linea == '':
            continue
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centrado else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.line_spacing = 1.0

        run = p.add_run(linea)
        run.font.name  = FUENTE
        run.font.size  = tamaño
        run.font.bold  = negrita
        run.font.color.rgb = RGBColor(
            int(color_texto[0:2], 16), int(color_texto[2:4], 16), int(color_texto[4:6], 16)
        )
        primer = False

def agregar_titulo_cuadro(doc, texto, bookmark_id=None):
    """
    Inserta el título de un cuadro ANTES de la tabla (Art. 6, 8, 13 IUTECP).
    - Título arriba del cuadro.
    - El número del cuadro va en negrita, la descripción en cursiva (Art. 6).
    - Espacio doble antes y después (Art. 8).
    Formato esperado del texto: 'Cuadro X. Descripción del cuadro'
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Espacio doble antes y después de cuadros titulados (Art. 8)
    p.paragraph_format.space_before = ESP_DOBLE
    p.paragraph_format.space_after = ESP_SENCILLO
    p.paragraph_format.line_spacing = INTERLINEADO
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True

    # Separar 'Cuadro X' de la descripción para aplicar formatos distintos
    partes = texto.split('. ', 1)
    if len(partes) == 2:
        # "Cuadro X" en negrita
        run_num = p.add_run(partes[0] + '. ')
        run_num.font.name = FUENTE
        run_num.font.size = Pt(TAMANO_BASE)
        run_num.font.bold = True
        # Descripción en cursiva (Art. 6)
        run_desc = p.add_run(partes[1])
        run_desc.font.name = FUENTE
        run_desc.font.size = Pt(TAMANO_BASE)
        run_desc.font.italic = True
    else:
        run = p.add_run(texto)
        run.font.name = FUENTE
        run.font.size = Pt(TAMANO_BASE)
        run.font.bold = True
    if bookmark_id:
        _agregar_bookmark(p, bookmark_id)

def _mantener_tabla_junta(tabla):
    """Evita que las filas de una tabla se partan entre páginas y hace que la
    tabla completa se mueva a la siguiente página si no cabe entera en la actual.

    - w:cantSplit en cada fila: la fila nunca se parte a mitad de contenido.
    - keep_with_next en los párrafos de cada celda (menos la última fila):
      "pega" las filas entre sí para que la tabla baje entera y no quede
      recortada al final de la página.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    rows = list(tabla.rows)
    for fila in rows:
        trPr = fila._tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        trPr.append(cant_split)
    for fila in rows[:-1]:
        for cell in fila.cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_with_next = True

def agregar_tabla_planificacion(doc, datos, titulo_cuadro=None, bookmark_id=None):
    """Genera la tabla de planificación de objetivos con 5 columnas"""
    if titulo_cuadro:
        agregar_titulo_cuadro(doc, titulo_cuadro, bookmark_id=bookmark_id)

    ENCABEZADOS = ['Objetivo', 'Variable', 'Actividades', 'Técnica', 'Instrumento']
    ANCHOS_CM   = [2.9, 2.9, 2.9, 2.9, 2.99]

    tabla = doc.add_table(rows=1 + len(datos), cols=5)
    tabla.style = 'Table Grid'
    aplicar_formato_tabla_xml(tabla, ANCHOS_CM)

    for col, enc in enumerate(ENCABEZADOS):
        cell = tabla.cell(0, col)
        set_cell_shading(cell, COLOR_ENCABEZADO)
        _celda(cell, enc, negrita=True, centrado=True, tamaño=Pt(TAMANO_TABLA), color_texto=COLOR_TEXTO_CLARO)

    for fila, datos_fila in enumerate(datos, start=1):
        if len(datos_fila) == 5:
            contenidos = list(datos_fila)
        else:
            contenidos = list(datos_fila) + ['']
        fondo = COLOR_FILA_PAR if fila % 2 == 0 else COLOR_TEXTO_CLARO
        for col, texto in enumerate(contenidos):
            cell = tabla.cell(fila, col)
            set_cell_shading(cell, fondo)
            _celda(cell, texto, tamaño=Pt(TAMANO_TABLA))

    _mantener_tabla_junta(tabla)

def agregar_gantt(doc, semanas, titulo_cuadro=None, bookmark_id=None):
    """Genera la tabla Gantt con 2 filas de encabezado (meses y semanas).
    Cada fila puede ser un dict {'desc', 'semanas', 'tipo'} donde `tipo` es
    'operativa', 'investigacion' o 'mixta' (define el color de la barra),
    o una tupla (desc, activas) para compatibilidad."""
    if titulo_cuadro:
        agregar_titulo_cuadro(doc, titulo_cuadro, bookmark_id=bookmark_id)

    if semanas and isinstance(semanas[0], dict):
        num_sem = max(len(item.get('semanas', [])) for item in semanas)
    else:
        num_sem = max(len(s[1]) for s in semanas)

    section = doc.sections[-1]
    ancho_util_emu = section.page_width - section.left_margin - section.right_margin
    ancho_util_cm = Emu(ancho_util_emu).cm

    COL_SEM_CM = Pt(3 * TAMANO_TABLA_CHICO * 0.55).cm + 0.4
    min_act_cm = 3.0
    max_s_por_col = (ancho_util_cm - min_act_cm) / num_sem if num_sem > 0 else ancho_util_cm
    COL_SEM_CM = min(COL_SEM_CM, max_s_por_col)
    COL_ACT_CM = ancho_util_cm - num_sem * COL_SEM_CM

    tabla = doc.add_table(rows=2 + len(semanas), cols=1 + num_sem)
    tabla.style = 'Table Grid'

    anchos_gantt = [COL_ACT_CM] + [COL_SEM_CM] * num_sem
    aplicar_formato_tabla_xml(tabla, anchos_gantt)

    # ── Fila 0: meses (JUNIO, JULIO, AGOSTO) ──
    cell_label = tabla.cell(0, 0)
    set_cell_shading(cell_label, COLOR_ENCABEZADO)
    _celda(cell_label, 'Semana', negrita=True, centrado=True, tamaño=Pt(TAMANO_TABLA), color_texto=COLOR_TEXTO_CLARO)

    meses = [("JUNIO", 4), ("JULIO", 4), ("AGOSTO", 2)]
    col_start = 1
    for nombre, ncols in meses:
        if col_start > num_sem:
            break
        ncols = min(ncols, num_sem - col_start + 1)
        c1 = tabla.cell(0, col_start)
        c2 = tabla.cell(0, col_start + ncols - 1)
        merged = c1.merge(c2)
        set_cell_shading(merged, COLOR_ENCABEZADO)
        _celda(merged, nombre, negrita=True, centrado=True, tamaño=Pt(TAMANO_TABLA), color_texto=COLOR_TEXTO_CLARO)
        col_start += ncols

    # ── Fila 1: semanas (1..n) ──
    cell_act = tabla.cell(1, 0)
    set_cell_shading(cell_act, COLOR_ENCABEZADO)
    _celda(cell_act, 'Actividad', negrita=True, centrado=True, tamaño=Pt(TAMANO_TABLA), color_texto=COLOR_TEXTO_CLARO)

    semana_meses = [4, 4, 2]  # weeks per month
    col_actual = 1
    for ncols in semana_meses:
        for w in range(1, ncols + 1):
            if col_actual > num_sem:
                break
            cell_s = tabla.cell(1, col_actual)
            set_cell_shading(cell_s, COLOR_ENCABEZADO)
            _celda(cell_s, str(w), negrita=True, centrado=True, tamaño=Pt(TAMANO_TABLA_CHICO), color_texto=COLOR_TEXTO_CLARO)
            col_actual += 1

    # ── Filas de datos ──
    colores_tipo = {
        'operativa': COLOR_GANTT_VERDE,
        'investigacion': COLOR_GANTT_INVESTIGACION,
        'proyecto': COLOR_GANTT_INVESTIGACION,
        'mixta': COLOR_GANTT_MIXTA,
    }
    for fila, item in enumerate(semanas, start=2):
        if isinstance(item, dict):
            desc = item.get('desc', '')
            activas = item.get('semanas', [])
            tipo = item.get('tipo')
        else:
            desc, activas = item
            tipo = None
        color_barra = colores_tipo.get(tipo, COLOR_GANTT_VERDE)
        fondo_fila = COLOR_FILA_PAR if fila % 2 == 0 else COLOR_TEXTO_CLARO
        cell_a = tabla.cell(fila, 0)
        set_cell_shading(cell_a, fondo_fila)
        _celda(cell_a, desc, tamaño=Pt(9))

        for s in range(num_sem):
            cell_s = tabla.cell(fila, s + 1)
            activa = activas[s] if s < len(activas) else False
            set_cell_shading(cell_s, color_barra if activa else fondo_fila)
            _celda(cell_s, '✓' if activa else '', centrado=True, tamaño=Pt(TAMANO_TABLA_CHICO), color_texto=COLOR_TEXTO_CLARO if activa else COLOR_TEXTO_OSCURO)

    _mantener_tabla_junta(tabla)
    _agregar_leyenda_gantt(doc, semanas)

def _agregar_leyenda_gantt(doc, semanas):
    """Leyenda de colores bajo la tabla Gantt (Trabajo operativo / proyecto / mixto)."""
    tipos_vistos = []
    for item in semanas:
        if isinstance(item, dict):
            tipo = item.get('tipo')
        else:
            tipo = None
        if tipo not in tipos_vistos:
            tipos_vistos.append(tipo)
    leyenda = {
        None: ('Trabajo operativo', COLOR_GANTT_VERDE),
        'operativa': ('Trabajo operativo', COLOR_GANTT_VERDE),
        'investigacion': ('Trabajo de proyecto / investigación', COLOR_GANTT_INVESTIGACION),
        'proyecto': ('Trabajo de proyecto / investigación', COLOR_GANTT_INVESTIGACION),
        'mixta': ('Trabajo operativo / proyecto', COLOR_GANTT_MIXTA),
    }
    legenda = [leyenda[t] for t in tipos_vistos if t in leyenda]
    if not legenda:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    for label, color in legenda:
        r = p.add_run('■ ')
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(color)
        r_lbl = p.add_run(label)
        r_lbl.font.name = FUENTE
        r_lbl.font.size = Pt(9)
        r_lbl.font.color.rgb = RGBColor.from_string(color)
        sep = p.add_run('        ')
        sep.font.size = Pt(9)

def _insertar_campo_pagina(run, formato_pagina='PAGE'):
    """Inserta un campo de número de página con el formato especificado en un run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = formato_pagina
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.extend([fldChar1, instrText, fldChar2, fldChar3])
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def agregar_numeracion_pie(doc, idx_inicio_cuerpo=None, idx_indice_inicio=None, idx_indice_fin=None):
    """
    Numeración de páginas según norma IUTECP:
    - Portada y contraportada (secciones 0 y 1): SIN número.
    - Preliminares (Aprobaciones, Agradecimientos, Dedicatoria, Índices, Listas,
      Resumen): Romanos en minúsculas (i, ii, iii...) con conteo continuo desde
      la portada. No se imprime en la primera página de los arranques que
      activan título de página (different_first_page_header_footer).
    - Cuerpo (Introducción, Capítulos I..V, Referencias, Anexos): Arábigos que
      REINICIAN en 1 en la Introducción. La primera página de la Introducción
      no lleva número, la segunda muestra "2", el Capítulo I arranca en "3"
      (sin número en su primera página) y continúa 4, 5, ...
    idx_inicio_cuerpo: índice de la sección donde empieza el cuerpo arábigo
    (normalmente la sección NEW_PAGE de la Introducción).
    """
    # 1. Desactivar numeración en portada y contraportada (secciones 0 y 1)
    for sec_idx in range(min(2, len(doc.sections))):
        section = doc.sections[sec_idx]
        footer = section.footer
        footer.is_linked_to_previous = False
        # Dejar el pie vacío

    # 2. Si no se proporcionó el índice de inicio del cuerpo, detectarlo
    #    buscando la primera sección cuyo primer párrafo contenga "CAPÍTULO"
    if idx_inicio_cuerpo is None:
        for i, sec in enumerate(doc.sections):
            # Buscar si algún párrafo del cuerpo de esta sección dice CAPÍTULO
            body_elem = sec._sectPr.getparent()
            # Simplemente iteramos las secciones y buscamos
            pass
        # Valor por defecto seguro: la Introducción es la última preliminar.
        # Las preliminares producen 2 secciones cada una (NEW_PAGE + CONTINUOUS).
        # Contamos: portada(0), contraportada(1), dedicatoria(2,3),
        # agradecimientos(4,5), resumen(6,7), introducción(8,9) => cap I empieza en 10
        idx_inicio_cuerpo = 10

    # 2b. La INTRODUCCIÓN reinicia la numeración arábiga en 1. Su primera página
    #     queda sin número (titlePg), por lo que la segunda página del índice
    #     físico pasa a ser la "2", el Capítulo I arranca en "3" (sin número) y
    #     las siguientes son "4", "5", ... Los capítulos, referencias y anexos
    #     continúan ese mismo conteo (sin volver a reiniciarse).
    if idx_inicio_cuerpo < len(doc.sections):
        sectPr = doc.sections[idx_inicio_cuerpo]._sectPr
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:start'), '1')
        # Insertar w:pgNumType en la posición correcta del esquema CT_SectPr
        # (debe ir antes de w:cols / w:titlePg / w:docGrid, tras w:pgMar).
        sectPr.insert(0, pgNumType)
        for child in list(sectPr):
            if child.tag == qn('w:cols') or child.tag == qn('w:titlePg') or child.tag == qn('w:docGrid'):
                sectPr.remove(pgNumType)
                child.addprevious(pgNumType)
                break

    # 3. Preliminares (incluidos los Índices) e Índice de Contenido: romanos en minúsculas
    #    (desde sección 2 hasta idx_inicio_cuerpo - 1). La primera página de cada
    #    sección que arranca con titlePg queda sin número (pie de primera página vacío).
    for sec_idx in range(2, len(doc.sections)):
        section = doc.sections[sec_idx]
        footer = section.footer
        footer.is_linked_to_previous = False
        formato = ' PAGE \\* roman ' if sec_idx < idx_inicio_cuerpo else ' PAGE '
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        _insertar_campo_pagina(run, formato)
        # Si la sección activa título de página (primera página sin número), asegurar
        # un pie de primera página vacío y desvinculado para que el número no aparezca allí.
        if section.different_first_page_header_footer:
            first = section.first_page_footer
            first.is_linked_to_previous = False

def buscar_imagen_por_numero(carpeta, numero, extensiones=None):
    """Busca una imagen por su nombre numérico en una carpeta específica"""
    if extensiones is None:
        extensiones = ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp']
    if not os.path.exists(carpeta):
        print(f"⚠ Advertencia: La carpeta '{carpeta}' no existe")
        return None
    numero_str = str(numero)
    for archivo in os.listdir(carpeta):
        nombre, ext = os.path.splitext(archivo)
        if nombre == numero_str and ext.lower()[1:] in extensiones:
            ruta_completa = os.path.join(carpeta, archivo)
            print(f"✓ Imagen encontrada: {ruta_completa}")
            return ruta_completa
    print(f"⚠ No se encontró imagen con número {numero} en {carpeta}")
    return None

def agregar_imagen(doc, ruta_imagen, titulo, ancho=Cm(12), fuente=None, bookmark_id=None):
    """
    Agrega una imagen (gráfico) y su título descriptivo DEBAJO (Art. 13 IUTECP).
    - Espacio doble antes y después de gráficos titulados (Art. 8).
    - El número del gráfico en cursiva, la descripción en negrita (Art. 6).
    - Opcionalmente agrega la línea 'Fuente:' (Art. 13).
    Formato esperado del titulo: 'Gráfico X. Descripción del gráfico'
    """
    if not ruta_imagen or not os.path.exists(ruta_imagen):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[IMAGEN NO ENCONTRADA: {titulo}]")
        run.font.italic = True
        run.font.color.rgb = RGBColor(255, 0, 0)
        return

    section = doc.sections[-1]
    max_width = Emu(section.page_width - section.left_margin - section.right_margin)
    if ancho > max_width:
        ancho = max_width

    # Espacio doble antes del gráfico (Art. 8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    run = p.add_run()
    try:
        run.add_picture(ruta_imagen, width=ancho)
    except Exception as e:
        print(f"❌ Error al agregar imagen {ruta_imagen}: {e}")
        return

    # Título del gráfico debajo de la imagen (Art. 13)
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_before = Pt(6)
    p_titulo.paragraph_format.space_after = Pt(24)
    p_titulo.paragraph_format.keep_together = True

    # Separar 'Gráfico X' de la descripción para formatos distintos
    partes = titulo.split('. ', 1)
    if len(partes) == 2:
        # "Gráfico X." en cursiva
        run_num = p_titulo.add_run(partes[0] + '. ')
        run_num.font.name = 'Times New Roman'
        run_num.font.size = Pt(12)
        run_num.font.italic = True
        # Descripción en negrita
        run_desc = p_titulo.add_run(partes[1])
        run_desc.font.name = 'Times New Roman'
        run_desc.font.size = Pt(12)
        run_desc.font.bold = True
    else:
        run_titulo = p_titulo.add_run(titulo)
        run_titulo.font.name = 'Times New Roman'
        run_titulo.font.size = Pt(12)
        run_titulo.font.bold = True
    if bookmark_id:
        _agregar_bookmark(p_titulo, bookmark_id)

    # Línea de Fuente opcional (Art. 13)
    if fuente:
        p_fuente = doc.add_paragraph()
        p_fuente.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_fuente.paragraph_format.space_before = Pt(0)
        p_fuente.paragraph_format.space_after = Pt(12)
        run_f_label = p_fuente.add_run('Fuente: ')
        run_f_label.font.name = 'Times New Roman'
        run_f_label.font.size = Pt(12)
        run_f_label.font.italic = True
        run_f_text = p_fuente.add_run(fuente)
        run_f_text.font.name = 'Times New Roman'
        run_f_text.font.size = Pt(12)

# ================================================================
# CONSTRUCCIÓN DEL DOCUMENTO (PORTADA COMPATIBLE LIBREOFFICE)
# ================================================================
def construir_portada(doc, solo_autor=False, idx_seccion=0):
    """
    Construye la portada distribuyendo los 4 bloques de forma proporcional
    al área útil de la página, sin valores fijos de puntos.

    Las posiciones de cada bloque se definen como porcentajes del área útil:
      POS_TITULO: dónde empieza el título (defecto 42% = ligeramente sobre el centro)
      POS_AUTOR : dónde empiezan los datos del autor (defecto 67% = tercio inferior)
      POS_FECHA : dónde empieza la fecha (defecto 90% = cerca del margen inferior)

    solo_autor: si es True, solo muestra el nombre y CI del autor (sin tutores),
                usado en la primera página (portada). La contraportada muestra todo.
    """
    # ------------------------------------------------------------------
    # 1. Leer dimensiones reales de la página (en puntos, 1 pt = 1/72 in)
    # ------------------------------------------------------------------
    section = doc.sections[idx_seccion]
    EMU_PER_PT = 12700  # 1 pt = 12700 EMU

    usable_h = (section.page_height - section.top_margin - section.bottom_margin) / EMU_PER_PT
    usable_w = (section.page_width  - section.left_margin - section.right_margin ) / EMU_PER_PT

    # ------------------------------------------------------------------
    # 2. Estimar la altura de cada bloque en puntos
    #    (fuente 12pt × interlineado 1.15 = ~13.8pt por línea)
    # ------------------------------------------------------------------
    LINE_PT = 12 * 1.15

    membrete_lines = len(c.MEMBRETE)

    lineas_autor_completo = c.AUTOR_DATOS if isinstance(c.AUTOR_DATOS, list) else [c.AUTOR_DATOS]
    if solo_autor:
        # Portada: solo nombre y CI (todo antes del primer string vacío)
        lineas_autor = []
        for linea in lineas_autor_completo:
            if not linea.strip():
                break
            lineas_autor.append(linea)
    else:
        lineas_autor = lineas_autor_completo
    autor_lines  = sum(1 for l in lineas_autor if l.strip())

    # Times New Roman 12pt: ancho promedio ~6pt por carácter
    chars_per_line = max(1, int(usable_w / 6.0))
    titulo_lines   = max(1, math.ceil(len(c.TITULO_PROYECTO) / chars_per_line))

    h_membrete = membrete_lines * LINE_PT
    h_titulo   = titulo_lines   * LINE_PT
    h_autor    = autor_lines    * LINE_PT
    h_fecha    = LINE_PT

    # ------------------------------------------------------------------
    # 3. Calcular gaps en puntos entre bloques
    #    El gap datos-fecha es la mitad del gap normal
    # ------------------------------------------------------------------
    h_membrete = membrete_lines * LINE_PT
    h_titulo   = titulo_lines   * LINE_PT
    h_autor    = autor_lines    * LINE_PT
    h_fecha    = LINE_PT

    # Logo height estimate (solo portada, 3cm de ancho ≈ 80pt de alto)
    logo_path = os.path.join("compartido", "iutecp.png")
    tiene_logo = solo_autor and os.path.exists(logo_path)
    h_logo = 67.0 if tiene_logo else 0

    h_contenido = h_membrete + h_titulo + h_autor + h_fecha
    gap = max(10.0, (usable_h - h_contenido - 30) / 3.1)
    gap_mt = gap * 1.1   # membrete-título más grande
    gap_td = gap          # título-datos
    gap_df = max(4.0, gap * 0.35)  # datos-fecha más chico que la mitad

    if solo_autor:
        # Portada: el logo se intercala entre el membrete y el título.
        if tiene_logo:
            logo_mitad = max(0, gap_mt / 4.0)
            logo_resto = max(0, gap_mt - logo_mitad - h_logo)
            before_titulo = 0
        else:
            logo_mitad = 0
            logo_resto = 0
            before_titulo = gap_mt
        before_autor  = gap_td
        before_fecha  = gap_df

    # La contraportada necesita una distribución más compacta porque incluye
    # la descripción institucional y los datos de ambos tutores. Los valores se
    # derivan del área útil para que la fecha quede bien ubicada sin desbordarse.
    if not solo_autor:
        before_titulo = max(40, min(64, int(usable_h * 0.10)))
        before_autor  = max(60, min(105, int(usable_h * 0.16)))
        before_fecha  = max(20, min(40, int(usable_h * 0.06)))

    # ------------------------------------------------------------------
    # 4. Renderizar cada bloque
    # ------------------------------------------------------------------

    # BLOQUE 1: MEMBRETE (Alineado al margen superior)
    p_memb = doc.add_paragraph()
    p_memb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_memb.paragraph_format.line_spacing = 1.0 if not solo_autor else 1.15
    p_memb.paragraph_format.space_before = Pt(0)
    p_memb.paragraph_format.space_after  = Pt(0)
    for i, linea in enumerate(c.MEMBRETE):
        r = p_memb.add_run(linea)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.bold = True
        if i < len(c.MEMBRETE) - 1:
            r.add_break()

    # BLOQUE 1.5: LOGO IUTECP (solo portada, debajo del membrete)
    if solo_autor and os.path.exists(logo_path):
        pic = doc.add_picture(logo_path, width=Cm(2.4))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        last_p.paragraph_format.space_before = Pt(logo_mitad)
        last_p.paragraph_format.space_after  = Pt(logo_resto)

    # BLOQUE 2: TÍTULO DEL PROYECTO (Centrado proporcionalmente)
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.line_spacing = 1.15
    p_titulo.paragraph_format.space_before = Pt(before_titulo)
    p_titulo.paragraph_format.space_after  = Pt(0)
    r_title = p_titulo.add_run(c.TITULO_PROYECTO)
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(12)
    r_title.font.bold = True

    if not solo_autor:
        p_desc = doc.add_paragraph()
        p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_desc.paragraph_format.line_spacing = 1.15
        p_desc.paragraph_format.space_before = Pt(22)
        p_desc.paragraph_format.space_after = Pt(0)
        r_desc = p_desc.add_run(
            "Informe de pasantías para obtener el título de Técnico Superior Universitario "
            f"en la especialidad de: {getattr(c, 'ESPECIALIDAD', '')}"
        )
        r_desc.font.name = 'Times New Roman'
        r_desc.font.size = Pt(12)
        r_desc.font.bold = True

    # BLOQUE 3: DATOS DEL AUTOR Y TUTORES
    if solo_autor:
        # Portada: todo el bloque alineado a la derecha
        p_datos = doc.add_paragraph()
        p_datos.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_datos.paragraph_format.line_spacing = 1.15
        p_datos.paragraph_format.space_before = Pt(before_autor)
        p_datos.paragraph_format.space_after  = Pt(0)
        for i, linea in enumerate(lineas_autor):
            r = p_datos.add_run(linea)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            r.font.bold = True
            if i < len(lineas_autor) - 1:
                r.add_break()
    else:
        # Contraportada: tutores a la izquierda, autor a la derecha, misma fila
        split_idx = len(lineas_autor)
        for i, linea in enumerate(lineas_autor):
            if not linea.strip():
                split_idx = i
                break
        autor_block = lineas_autor[:split_idx]
        tutor_start = split_idx
        while tutor_start < len(lineas_autor) and not lineas_autor[tutor_start].strip():
            tutor_start += 1
        tutor_block = lineas_autor[tutor_start:]

        # Espaciador antes de la tabla (para posicionar a la altura de POS_AUTOR)
        p_before = doc.add_paragraph()
        p_before.paragraph_format.space_before = Pt(before_autor)
        p_before.paragraph_format.space_after = Pt(0)
        p_before.paragraph_format.line_spacing = Pt(1)

        # Tabla invisible de 1 fila × 2 columnas
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.allow_autofit = False

        # Quitar bordes
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            borders.append(border)
        tblPr.append(borders)

        # Cell 0: tutores (izquierda)
        cell_tutor = table.rows[0].cells[0]
        cell_tutor.width = Cm(7.29)
        cell_tutor.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p_t = cell_tutor.paragraphs[0]
        p_t.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_t.paragraph_format.line_spacing = 1.0
        for i, linea in enumerate(tutor_block):
            run = p_t.add_run(linea)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            if i < len(tutor_block) - 1:
                run.add_break()

        # Cell 1: autor (derecha)
        cell_autor = table.rows[0].cells[1]
        cell_autor.width = Cm(7.30)
        cell_autor.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p_a = cell_autor.paragraphs[0]
        p_a.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_a.paragraph_format.line_spacing = 1.0
        for i, linea in enumerate(autor_block):
            run = p_a.add_run(linea)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            if i < len(autor_block) - 1:
                run.add_break()

    # BLOQUE 4: CIUDAD Y FECHA (Proporcional al margen inferior)
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pie.paragraph_format.line_spacing = 1.15
    p_pie.paragraph_format.space_before = Pt(before_fecha)
    p_pie.paragraph_format.space_after  = Pt(0)
    r_pie = p_pie.add_run(c.FECHA_LUGAR)
    r_pie.font.name = 'Times New Roman'
    r_pie.font.size = Pt(12)
    r_pie.font.bold = True

def agregar_parada_tabulacion_puntos(parrafo, posicion_emu):
    """Agrega una parada de tabulación derecha con puntos de relleno usando la API oficial de python-docx."""
    parrafo.paragraph_format.tab_stops.add_tab_stop(
        Emu(posicion_emu),
        alignment=WD_TAB_ALIGNMENT.RIGHT,
        leader=WD_TAB_LEADER.DOTS
    )

def agregar_tabulaciones_lista(parrafo, pos_tab_texto_emu, pos_tab_pag_emu):
    """Agrega dos paradas de tabulación oficiales: izquierda (texto) y derecha con puntos (página)."""
    parrafo.paragraph_format.tab_stops.add_tab_stop(
        Emu(pos_tab_texto_emu),
        alignment=WD_TAB_ALIGNMENT.LEFT
    )
    parrafo.paragraph_format.tab_stops.add_tab_stop(
        Emu(pos_tab_pag_emu),
        alignment=WD_TAB_ALIGNMENT.RIGHT,
        leader=WD_TAB_LEADER.DOTS
    )

def agregar_tabulacion_derecha(parrafo, pos_tab_pag_emu):
    """Agrega una única parada de tabulación derecha oficial (para alinear pp. en la cabecera)."""
    parrafo.paragraph_format.tab_stops.add_tab_stop(
        Emu(pos_tab_pag_emu),
        alignment=WD_TAB_ALIGNMENT.RIGHT
    )

def agregar_fila_indice_general_nativa(doc, titulo, pagina, sangria_cm=0, negrita=False, bookmark_id=None, valor_calculado=''):
    """Agrega una línea del índice general usando tabulaciones nativas de Word para alinear al extremo derecho de forma absoluta.

    Si se pasa `bookmark_id`, en lugar del número `pagina` se inserta un campo
    PAGEREF que Word/LibreOffice rellena automáticamente con la página real
    del título marcado por ese bookmark. `valor_calculado` es el texto visible
    mientras el campo no se actualiza.
    """
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(sangria_cm)
    
    section = doc.sections[-1]
    ancho_util_emu = section.page_width - section.left_margin - section.right_margin
    # La posición del tabulador en Word es absoluta respecto a los márgenes, por lo que usamos directamente ancho_util_emu
    agregar_parada_tabulacion_puntos(p, ancho_util_emu)
    
    run_desc = p.add_run(f"{titulo}\t")
    run_desc.font.name = 'Times New Roman'
    run_desc.font.size = Pt(12)
    run_desc.font.bold = negrita
    
    if bookmark_id:
        respaldo = valor_calculado if valor_calculado else pagina
        _agregar_campo_pageref(p, bookmark_id, negrita=negrita, valor_calculado=respaldo)
    else:
        run_pag = p.add_run(pagina)
        run_pag.font.name = 'Times New Roman'
        run_pag.font.size = Pt(12)
        run_pag.font.bold = negrita
    return p

def _agregar_bookmark(parrafo, bookmark_id):
    """Inserta un marcador (bookmark) XML oculto en el párrafo indicado.

    El bookmark rodea el contenido del párrafo, permitiendo que un campo
    PAGEREF del índice apunte a la página real donde cae este párrafo.
    """
    p_elem = parrafo._element
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), bookmark_id)
    start.set(qn('w:name'), bookmark_id)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), bookmark_id)
    # Insertar al inicio y al final del párrafo para rodear todo su contenido
    p_elem.insert(0, start)
    p_elem.append(end)

def _agregar_campo_pageref(parrafo, bookmark_id, negrita=False, valor_calculado='?'):
    """Inserta un campo PAGEREF en el párrafo que Word/LibreOffice evalúa
    al actualizar campos, mostrando la página real del bookmark indicado.

    `valor_calculado`: texto mostrado mientras el campo no se ha actualizado
    (respaldo visible en visores que no recalculan campos, p. ej. PDFs headless).
    """
    run = parrafo.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = negrita

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' PAGEREF {bookmark_id} \\h '

    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')

    # Texto de respaldo que se muestra si los campos no se actualizan
    t = OxmlElement('w:t')
    t.text = valor_calculado

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run._element.append(fldChar_begin)
    run._element.append(instrText)
    run._element.append(fldChar_sep)
    run._element.append(t)
    run._element.append(fldChar_end)

def agregar_fila_lista_preliminar_nativa(doc, col1_text, col2_text, col3_text, bookmark_id=None):
    """Agrega una entrada a una lista descriptiva (cuadro/gráfico/anexo) con sangría colgante y tabulación nativa absoluta.

    Si se pasa `bookmark_id`, la tercera columna (página) se reemplaza por un
    campo PAGEREF que Word/LibreOffice actualiza automáticamente.
    """
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.8)
    p.paragraph_format.first_line_indent = Cm(-1.8)

    section = doc.sections[-1]
    ancho_util_emu = section.page_width - section.left_margin - section.right_margin
    # La posición del tabulador en Word es absoluta respecto a los márgenes, por lo que usamos directamente ancho_util_emu
    agregar_tabulaciones_lista(p, Cm(1.8).emu, ancho_util_emu)

    # Escribir número y descripción con formato
    run = p.add_run(f"{col1_text}\t{col2_text}\t")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    if bookmark_id:
        _agregar_campo_pageref(p, bookmark_id, negrita=False, valor_calculado=col3_text)
    else:
        run_pag = p.add_run(col3_text)
        run_pag.font.name = 'Times New Roman'
        run_pag.font.size = Pt(12)
    return p

# ================================================================
#  INSERCIÓN DE GRÁFICOS DATA-DRIVEN
# ================================================================
# Anclas válidas donde pueden insertarse gráficos:
#   "ubicacion"  -> tras la sección 1.1.7 Ubicación geográfica
#   "estructura" -> tras la sección 1.1.9 Estructura Organizativa
ANCLAS_VALIDAS = {"ubicacion", "estructura"}

def _resolver_ruta_imagen(carpeta_imagenes, cfg):
    """Resuelve una imagen por nombre explícito (`archivo`) o por número, manteniendo compatibilidad."""
    archivo = cfg.get("archivo") if isinstance(cfg, dict) else None
    if archivo:
        ruta = os.path.join(carpeta_imagenes, archivo)
        if os.path.exists(ruta):
            return ruta
        print(f"⚠ No se encontró imagen '{archivo}' en {carpeta_imagenes}")
        return None
    numero = cfg.get("numero") if isinstance(cfg, dict) else None
    return buscar_imagen_por_numero(carpeta_imagenes, numero)

def _insertar_graficos_por_ancla(doc, carpeta_imagenes, ancla):
    """Inserta gráficos (GRAFICOS) y figuras (FIGURAS) de contenido.py marcados con `tras = ancla`."""
    graficos = getattr(c, 'GRAFICOS', [])
    for g in graficos:
        if g.get("tras") != ancla:
            continue
        numero = g.get("numero")
        titulo = g.get("titulo", f"Gráfico {numero}.")
        ancho = g.get("ancho_cm", 12)
        ruta = _resolver_ruta_imagen(carpeta_imagenes, g)
        bookmark_id = f"bm_grafico{numero}" if numero else None
        agregar_imagen(doc, ruta, titulo, ancho=Cm(ancho), fuente=g.get("fuente"), bookmark_id=bookmark_id)
    figuras = getattr(c, 'FIGURAS', [])
    for f in figuras:
        if f.get("tras") != ancla:
            continue
        numero = f.get("numero")
        titulo = f.get("titulo", f"Figura {numero}.")
        ancho = f.get("ancho_cm", 15)
        ruta = _resolver_ruta_imagen(carpeta_imagenes, f)
        bookmark_id = f"bm_figura{numero}" if numero else None
        if f.get("pagina_propia", False):
            doc.add_page_break()
        agregar_imagen(doc, ruta, titulo, ancho=Cm(ancho), fuente=f.get("fuente"), bookmark_id=bookmark_id)

def _insertar_logo_empresa(doc, carpeta_imagenes):
    """Inserta el logotipo configurado en GRAFICOS o mediante LOGO_EMPRESA."""
    logo_grafico = next(
        (g for g in getattr(c, 'GRAFICOS', []) if g.get("tras") == "logo_empresa"),
        None,
    )
    if logo_grafico:
        numero = logo_grafico.get("numero")
        ruta = _resolver_ruta_imagen(carpeta_imagenes, logo_grafico)
        agregar_imagen(
            doc,
            ruta,
            logo_grafico.get("titulo", f"Gráfico {numero}."),
            ancho=Cm(logo_grafico.get("ancho_cm", 12)),
            fuente=logo_grafico.get("fuente"),
            bookmark_id=f"bm_grafico{numero}" if numero else None,
        )
        return

    cfg = getattr(c, 'LOGO_EMPRESA', None)
    if not cfg:
        return
    if isinstance(cfg, str):
        cfg = {"archivo": cfg}
    ruta = _resolver_ruta_imagen(carpeta_imagenes, cfg)
    if not ruta:
        return
    ancho = Cm(cfg.get("ancho_cm", 4.0))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    try:
        p.add_run().add_picture(ruta, width=ancho)
    except Exception as e:
        print(f"❌ Error al agregar logo empresarial {ruta}: {e}")

def agregar_pagina_aprobacion(doc, titulo, texto_parrafo, pie_firma, nombre_tutor="", ci_tutor="", solo_encabezado=False):
    """Agrega página de aprobación con membrete, título centrado, firma y datos del tutor.

    Si solo_encabezado=True se deja la página en blanco tras el membrete y el título
    (la revisión de los tutores aún no está disponible).
    """
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _config_seccion(sec, primera_sin_numero=False)
    
    for linea in c.MEMBRETE:
        p_m = doc.add_paragraph()
        p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_m.paragraph_format.line_spacing = 1.15
        p_m.paragraph_format.space_before = Pt(0)
        p_m.paragraph_format.space_after = Pt(0)
        run_m = p_m.add_run(linea)
        run_m.font.name = FUENTE
        run_m.font.size = Pt(12)
        run_m.font.bold = True
    
    doc.add_paragraph()
    
    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t.paragraph_format.space_before = Pt(24)
    p_t.paragraph_format.space_after = Pt(24)
    run_t = p_t.add_run(titulo)
    run_t.font.name = FUENTE
    run_t.font.size = Pt(12)
    run_t.font.bold = True
    bookmark_id = "bm_aprob_ind" if "INDUSTRIAL" in titulo else "bm_aprob_acad"
    _agregar_bookmark(p_t, bookmark_id)

    if solo_encabezado:
        return
    
    p_ap = doc.add_paragraph()
    p_ap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_ap.paragraph_format.line_spacing = 1.5
    p_ap.paragraph_format.first_line_indent = Cm(0)
    run_ap = p_ap.add_run(texto_parrafo)
    run_ap.font.name = FUENTE
    run_ap.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p_f = doc.add_paragraph()
    p_f.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_f.paragraph_format.line_spacing = 1.5
    run_f = p_f.add_run(pie_firma)
    run_f.font.name = FUENTE
    run_f.font.size = Pt(12)
    
    for _ in range(1):
        doc.add_paragraph()
    
    # Línea de firma (o imagen de firma si está configurada en contenido.py)
    if "INDUSTRIAL" in titulo:
        firma_img = getattr(c, 'FIRMA_TUTOR_INDUSTRIAL', None)
    else:
        firma_img = getattr(c, 'FIRMA_TUTOR_ACADEMICO', None)

    ruta_firma = None
    if firma_img:
        ruta_firma = os.path.join(getattr(c, 'CARPETA_IMAGENES', 'imagenes'), firma_img)
        if not os.path.exists(ruta_firma):
            ruta_firma = None

    if ruta_firma:
        p_firma = doc.add_paragraph()
        p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_firma = p_firma.add_run()
        run_firma.add_picture(ruta_firma, width=Cm(4))
    else:
        p_linea = doc.add_paragraph()
        p_linea.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_linea = p_linea.add_run("_" * 35)
        run_linea.font.name = FUENTE
        run_linea.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Nombre del tutor centrado
    p_nom = doc.add_paragraph()
    p_nom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nom.paragraph_format.line_spacing = 1.5
    run_nom = p_nom.add_run(nombre_tutor)
    run_nom.font.name = FUENTE
    run_nom.font.size = Pt(12)
    run_nom.font.bold = True
    
    # Cédula del tutor centrada
    p_ci = doc.add_paragraph()
    p_ci.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ci.paragraph_format.space_before = Pt(0)
    run_ci = p_ci.add_run(ci_tutor)
    run_ci.font.name = FUENTE
    run_ci.font.size = Pt(12)

def construir_cuerpo_documento(doc, modo="completo"):
    """Escribe secuencialmente todas las secciones del informe de pasantía.
    
    modo: "completo" | "borrador1" (solo Cap I) | "borrador2" (Cap I+II) |
          "borrador3" (Cap I+II+III) | "borrador4" (Cap I+II+III+IV)
    """
    tiene_cap2 = modo in ("completo", "borrador2", "borrador3", "borrador4")
    tiene_cap3 = modo in ("completo", "borrador3", "borrador4")
    tiene_cap4 = modo in ("completo", "borrador4")
    tiene_cap5 = modo == "completo"

    # --- CÁLCULO DINÁMICO DE PÁGINAS PRELIMINARES ---
    pag_actual_romana = 5  # portada=i, contraportada=ii, aprob_tutor_industrial=iii, aprob_tutor_academico=iv
    romanos = {1: 'i', 2: 'ii', 3: 'iii', 4: 'iv', 5: 'v', 6: 'vi', 7: 'vii', 8: 'viii', 9: 'ix', 10: 'x', 11: 'xi', 12: 'xii'}
    
    pag_aprob_ind = romanos.get(3, 'iii')  # Siempre presente
    pag_aprob_acad = romanos.get(4, 'iv')  # Siempre presente
    pag_agradecimientos = ""
    pag_dedicatoria = ""
    pag_resumen = ""
    
    if hasattr(c, 'AGRADECIMIENTOS') and c.AGRADECIMIENTOS:
        pag_agradecimientos = romanos.get(pag_actual_romana, str(pag_actual_romana))
        pag_actual_romana += 1
        
    if hasattr(c, 'DEDICATORIA') and c.DEDICATORIA:
        pag_dedicatoria = romanos.get(pag_actual_romana, str(pag_actual_romana))
        pag_actual_romana += 1
        
    pag_indice = romanos.get(pag_actual_romana, str(pag_actual_romana))
    pag_actual_romana += 1
    
    pag_lista_cuadros = romanos.get(pag_actual_romana, str(pag_actual_romana))
    pag_actual_romana += 1

    pag_lista_figuras = ""
    if getattr(c, 'FIGURAS', []):
        pag_lista_figuras = romanos.get(pag_actual_romana, str(pag_actual_romana))
        pag_actual_romana += 1
    
    pag_lista_graficos = romanos.get(pag_actual_romana, str(pag_actual_romana))
    pag_actual_romana += 1
    
    pag_lista_anexos = ""
    if hasattr(c, 'ANEXOS_LISTA') and c.ANEXOS_LISTA:
        pag_lista_anexos = romanos.get(pag_actual_romana, str(pag_actual_romana))
        pag_actual_romana += 1

    if hasattr(c, 'RESUMEN_TEXTO') and c.RESUMEN_TEXTO:
        pag_resumen = romanos.get(pag_actual_romana, str(pag_actual_romana))
        pag_actual_romana += 1

    if True:
        # --- APROBACIÓN DEL TUTOR INDUSTRIAL ---
        nom_pas = getattr(c, 'NOMBRE_PASANTE', '[Nombre del Pasante]')
        ci_pas = getattr(c, 'CI_PASANTE', 'XX.XXX.XXX')
        esp = getattr(c, 'ESPECIALIDAD', '[Especialidad]')
        tit_proy = getattr(c, 'TITULO_PROYECTO', '[Título del Proyecto]')
        ciudad = getattr(c, 'CIUDAD_FECHA', 'El Tigre').split(",")[0] if "," in getattr(c, 'CIUDAD_FECHA', 'El Tigre') else 'El Tigre'
        
        texto_base = (
            f"En mi car\u00e1cter de tutor industrial del informe de pasant\u00edas presentado por: "
            f"{nom_pas}, de C\u00e9dula de Identidad {ci_pas}; para optar al grado de "
            f"T\u00e9cnico Superior Universitario en la especialidad de: {esp}, cuyo t\u00edtulo es; "
            f"\u201c{tit_proy}\u201d, manifiesto que cumple con los requisitos exigidos por el "
            f"Instituto Universitario de Tecnolog\u00eda \u201cEl\u00edas Calixto Pompa\u201d (IUTECP); "
            f"y que, por lo tanto, considero que re\u00fane los m\u00e9ritos suficientes para ser "
            f"evaluado por el jurado que se decida designar a tal fin."
        )
        # Extraer nombres y CIs de tutores desde AUTOR_DATOS
        autor_datos = getattr(c, 'AUTOR_DATOS', [])
        tut_ind_nom = ""
        tut_ind_ci = ""
        tut_acad_nom = ""
        tut_acad_ci = ""
        for i, linea in enumerate(autor_datos):
            if "Tutor Industrial" in linea and ":" in linea:
                partes = linea.split(":", 1)
                if len(partes) > 1 and partes[1].strip():
                    tut_ind_nom = partes[1].strip()
                elif i + 2 < len(autor_datos):
                    tut_ind_nom = autor_datos[i + 1].strip()
                    if "C.I.:" in autor_datos[i + 2]:
                        tut_ind_ci = autor_datos[i + 2].split(":", 1)[1].strip()
            if "Tutor Académico" in linea and ":" in linea:
                partes = linea.split(":", 1)
                if len(partes) > 1 and partes[1].strip():
                    tut_acad_nom = partes[1].strip()
                elif i + 2 < len(autor_datos):
                    tut_acad_nom = autor_datos[i + 1].strip()
                    if "C.I.:" in autor_datos[i + 2]:
                        tut_acad_ci = autor_datos[i + 2].split(":", 1)[1].strip()
            if "C.I.:" in linea and i > 0:
                prev = autor_datos[i - 1].strip()
                if prev == tut_ind_nom:
                    tut_ind_ci = linea.split(":", 1)[1].strip()
                elif prev == tut_acad_nom:
                    tut_acad_ci = linea.split(":", 1)[1].strip()

        agregar_pagina_aprobacion(doc, "APROBACIÓN DEL TUTOR INDUSTRIAL", texto_base,
            f"En la Ciudad de {ciudad}, a los XX días del mes de _______ de 2026",
            tut_ind_nom, f"C.I.: {tut_ind_ci}" if tut_ind_ci else "", solo_encabezado=False)

        # --- APROBACIÓN DEL TUTOR ACADÉMICO ---

        # --- APROBACIÓN DEL TUTOR ACADÉMICO ---
        texto_base_acad = (
            f"En mi car\u00e1cter de tutor acad\u00e9mico del informe de pasant\u00edas presentado por: "
            f"{nom_pas}, de C\u00e9dula de Identidad {ci_pas}; para optar al grado de "
            f"T\u00e9cnico Superior Universitario en la especialidad de: {esp}, cuyo t\u00edtulo es; "
            f"\u201c{tit_proy}\u201d, manifiesto que cumple con los requisitos exigidos por el "
            f"Instituto Universitario de Tecnolog\u00eda \u201cEl\u00edas Calixto Pompa\u201d (IUTECP); "
            f"y que, por lo tanto, considero que re\u00fane los m\u00e9ritos suficientes para ser "
            f"evaluado por el jurado que se decida designar a tal fin."
        )
        agregar_pagina_aprobacion(doc, "APROBACIÓN DEL TUTOR ACADÉMICO", texto_base_acad,
            f"En la Ciudad de {ciudad}, a los XX días del mes de _______ de 2026",
            tut_acad_nom, f"C.I.: {tut_acad_ci}" if tut_acad_ci else "", solo_encabezado=False)

        # --- PÁGINAS PRELIMINARES ---
        if hasattr(c, 'AGRADECIMIENTOS') and c.AGRADECIMIENTOS:
            iniciar_seccion_preliminar(doc, "AGRADECIMIENTOS", bookmark_id="bm_agradecimientos")
            agradecimientos = c.AGRADECIMIENTOS
            if isinstance(agradecimientos, (list, tuple)):
                for parrafo in agradecimientos:
                    agregar_parrafo_normado(doc, parrafo)
            else:
                agregar_parrafo_normado(doc, agradecimientos)
    
        if hasattr(c, 'DEDICATORIA') and c.DEDICATORIA:
            iniciar_seccion_preliminar(doc, "DEDICATORIA", bookmark_id="bm_dedicatoria")
            dedicatoria = c.DEDICATORIA
            if isinstance(dedicatoria, (list, tuple)):
                for parrafo in dedicatoria:
                    agregar_parrafo_normado(doc, parrafo)
            else:
                agregar_parrafo_normado(doc, dedicatoria)
    
        # --- ÍNDICE DE CONTENIDO ---
        idx_indice_inicio = len(doc.sections)
        iniciar_seccion_preliminar(doc, "ÍNDICE DE CONTENIDO")
        p_header_ind = doc.add_paragraph()
        p_header_ind.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_header_ind.paragraph_format.space_after = Pt(12)
        run_h_ind = p_header_ind.add_run("pp.")
        run_h_ind.font.name = 'Times New Roman'
        run_h_ind.font.size = Pt(12)
        run_h_ind.font.bold = True
    
        if pag_aprob_ind:
            agregar_fila_indice_general_nativa(doc, "APROBACIÓN DEL TUTOR INDUSTRIAL", "", bookmark_id="bm_aprob_ind", valor_calculado=pag_aprob_ind)
        if pag_aprob_acad:
            agregar_fila_indice_general_nativa(doc, "APROBACIÓN DEL TUTOR ACADÉMICO", "", bookmark_id="bm_aprob_acad", valor_calculado=pag_aprob_acad)
        if pag_agradecimientos:
            agregar_fila_indice_general_nativa(doc, "AGRADECIMIENTOS", "", bookmark_id="bm_agradecimientos", valor_calculado=pag_agradecimientos)
        if pag_dedicatoria:
            agregar_fila_indice_general_nativa(doc, "DEDICATORIA", "", bookmark_id="bm_dedicatoria", valor_calculado=pag_dedicatoria)
    
        agregar_fila_indice_general_nativa(doc, "LISTA DE CUADROS", "", bookmark_id="bm_lista_cuadros", valor_calculado=pag_lista_cuadros)
        if pag_lista_figuras:
            agregar_fila_indice_general_nativa(doc, "LISTA DE FIGURAS", "", bookmark_id="bm_lista_figuras", valor_calculado=pag_lista_figuras)
        agregar_fila_indice_general_nativa(doc, "LISTA DE GRÁFICOS", "", bookmark_id="bm_lista_graficos", valor_calculado=pag_lista_graficos)
        if pag_lista_anexos:
            agregar_fila_indice_general_nativa(doc, "LISTA DE ANEXOS", "", bookmark_id="bm_lista_anexos", valor_calculado=pag_lista_anexos)
        if pag_resumen:
            agregar_fila_indice_general_nativa(doc, "RESUMEN", "", bookmark_id="bm_resumen", valor_calculado=pag_resumen)
    
        agregar_fila_indice_general_nativa(doc, "INTRODUCCIÓN", "", bookmark_id="bm_introduccion", valor_calculado="1")

        idx_indice_fin = len(doc.sections)

        # Capítulos
        p_cap_lbl = doc.add_paragraph()
        p_cap_lbl.paragraph_format.space_before = Pt(12)
        p_cap_lbl.paragraph_format.space_after = Pt(6)
        p_cap_lbl.add_run("CAPÍTULOS").font.bold = True
    
        # Capítulo I
        agregar_fila_indice_general_nativa(doc, "CAPÍTULO I: REALIDAD ORGANIZACIONAL", "", sangria_cm=0, negrita=True, bookmark_id="bm_cap1")
        agregar_fila_indice_general_nativa(doc, "Identificación de la empresa", "", sangria_cm=0.5, bookmark_id="bm_cap1_ident")
        agregar_fila_indice_general_nativa(doc, "Reseña histórica", "", sangria_cm=0.5, bookmark_id="bm_cap1_resena")
        agregar_fila_indice_general_nativa(doc, "Misión", "", sangria_cm=0.5, bookmark_id="bm_cap1_mision")
        agregar_fila_indice_general_nativa(doc, "Visión", "", sangria_cm=0.5, bookmark_id="bm_cap1_vision")
        agregar_fila_indice_general_nativa(doc, "Valores", "", sangria_cm=0.5, bookmark_id="bm_cap1_valores")
        agregar_fila_indice_general_nativa(doc, "Objetivos Organizacionales", "", sangria_cm=0.5, bookmark_id="bm_cap1_obj")
        agregar_fila_indice_general_nativa(doc, "Ubicación geográfica", "", sangria_cm=0.5, bookmark_id="bm_cap1_ubic")
        agregar_fila_indice_general_nativa(doc, "Estructura organizacional de la empresa (organigrama)", "", sangria_cm=0.5, bookmark_id="bm_cap1_estruct")
    
        # Capítulo II
        if tiene_cap2:
            agregar_fila_indice_general_nativa(doc, "CAPÍTULO II: DIAGNÓSTICO SITUACIONAL", "", sangria_cm=0, negrita=True, bookmark_id="bm_cap2")
            agregar_fila_indice_general_nativa(doc, "Identificación de la situación problemática", "", sangria_cm=0.5, bookmark_id="bm_cap2_sit")
            agregar_fila_indice_general_nativa(doc, "Objetivo General", "", sangria_cm=0.5, bookmark_id="bm_cap2_objg")
            agregar_fila_indice_general_nativa(doc, "Objetivos Específicos", "", sangria_cm=0.5, bookmark_id="bm_cap2_obje")
            agregar_fila_indice_general_nativa(doc, "Planificación integral de objetivos", "", sangria_cm=0.5, bookmark_id="bm_cap2_planif")
            agregar_fila_indice_general_nativa(doc, "Cronograma de actividades", "", sangria_cm=0.5, bookmark_id="bm_cap2_crono")
    
        # Capítulo III
        if tiene_cap3:
            agregar_fila_indice_general_nativa(doc, "CAPÍTULO III: MARCO TEÓRICO", "", sangria_cm=0, negrita=True, bookmark_id="bm_cap3")
            agregar_fila_indice_general_nativa(doc, "Bases teóricas referenciales", "", sangria_cm=0.5, bookmark_id="bm_cap3_bases")
    
        # Capítulo IV
        if tiene_cap4:
            agregar_fila_indice_general_nativa(doc, "CAPÍTULO IV: ACTIVIDADES REALIZADAS", "", sangria_cm=0, negrita=True, bookmark_id="bm_cap4")
            agregar_fila_indice_general_nativa(doc, "Descripción de actividades ejecutadas por semana", "", sangria_cm=0.5, bookmark_id="bm_cap4_desc")
            actividades_indice = getattr(c, 'ACTIVIDADES_LISTA', [])
            for i, _act in enumerate(actividades_indice, 1):
                n_semana = _act.get('semana', i) if isinstance(_act, dict) else i
                agregar_fila_indice_general_nativa(doc, f"Semana {n_semana}", "", sangria_cm=1.0, bookmark_id=f"bm_cap4_s{i}")
    
        # Capítulo V
        if tiene_cap5:
            agregar_fila_indice_general_nativa(doc, "CAPÍTULO V: CONCLUSIONES Y RECOMENDACIONES", "", sangria_cm=0, negrita=True, bookmark_id="bm_cap5")
            agregar_fila_indice_general_nativa(doc, "Conclusiones", "", sangria_cm=0.5, bookmark_id="bm_cap5_concl")
            agregar_fila_indice_general_nativa(doc, "Recomendaciones", "", sangria_cm=0.5, bookmark_id="bm_cap5_recom")
    
        # Referencias y Anexos
        agregar_fila_indice_general_nativa(doc, "REFERENCIAS", "", sangria_cm=0, negrita=True, bookmark_id="bm_referencias")
        agregar_fila_indice_general_nativa(doc, "ANEXOS", "", sangria_cm=0, negrita=True, bookmark_id="bm_anexos")
    
        # --- LISTA DE CUADROS ---
        iniciar_seccion_preliminar(doc, "LISTA DE CUADROS", bookmark_id="bm_lista_cuadros")
        p_header_c = doc.add_paragraph()
        p_header_c.paragraph_format.space_before = Pt(12)
        p_header_c.paragraph_format.space_after = Pt(12)
    
        section_c = doc.sections[-1]
        ancho_c_emu = section_c.page_width - section_c.left_margin - section_c.right_margin
        agregar_tabulacion_derecha(p_header_c, ancho_c_emu)
    
        run_c_lbl = p_header_c.add_run("CUADRO")
        run_c_lbl.font.name = 'Times New Roman'
        run_c_lbl.font.size = Pt(12)
        run_c_lbl.font.bold = True
    
        p_header_c.add_run("\t")  # Salto de tabulador explícito para OXML
    
        run_pp_c = p_header_c.add_run("pp.")
        run_pp_c.font.name = FUENTE
        run_pp_c.font.size = Pt(TAMANO_BASE)
        run_pp_c.font.bold = True
    
        # Lista de Cuadros: data-driven desde contenido.py (CUADROS_INDICE) o defaults
        cuadros_indice_def = [
            ("1", "Planificación integral de objetivos específicos", "5"),
            ("2", "Cronograma de actividades administrativas", "6"),
        ]
        cuadros_indice = getattr(c, 'CUADROS_INDICE', cuadros_indice_def)
        if tiene_cap2:
            for num, desc, pag in cuadros_indice:
                bookmark_id = f"bm_cuadro{num}" if num else None
                agregar_fila_lista_preliminar_nativa(doc, num, desc, pag, bookmark_id=bookmark_id)
    
        # --- LISTA DE FIGURAS ---
        figuras_cfg = getattr(c, 'FIGURAS', [])
        if figuras_cfg:
            iniciar_seccion_preliminar(doc, "LISTA DE FIGURAS", bookmark_id="bm_lista_figuras")
            p_header_f = doc.add_paragraph()
            p_header_f.paragraph_format.space_before = Pt(12)
            p_header_f.paragraph_format.space_after = Pt(12)

            section_f = doc.sections[-1]
            ancho_f_emu = section_f.page_width - section_f.left_margin - section_f.right_margin
            agregar_tabulacion_derecha(p_header_f, ancho_f_emu)

            run_f_lbl = p_header_f.add_run("FIGURA")
            run_f_lbl.font.name = FUENTE
            run_f_lbl.font.size = Pt(TAMANO_BASE)
            run_f_lbl.font.bold = True
            p_header_f.add_run("\t")
            run_pp_f = p_header_f.add_run("pp.")
            run_pp_f.font.name = FUENTE
            run_pp_f.font.size = Pt(TAMANO_BASE)
            run_pp_f.font.bold = True

            for figura in figuras_cfg:
                num = str(figura.get("numero", ""))
                desc = figura.get("lista", figura.get("titulo", "").split('. ', 1)[-1])
                pag = str(figura.get("pagina", ""))
                bookmark_id = f"bm_figura{num}" if num else None
                agregar_fila_lista_preliminar_nativa(doc, num, desc, pag, bookmark_id=bookmark_id)

        # --- LISTA DE GRÁFICOS ---
        iniciar_seccion_preliminar(doc, "LISTA DE GRÁFICOS", bookmark_id="bm_lista_graficos")
        p_header_g = doc.add_paragraph()
        p_header_g.paragraph_format.space_before = Pt(12)
        p_header_g.paragraph_format.space_after = Pt(12)
    
        section_g = doc.sections[-1]
        ancho_g_emu = section_g.page_width - section_g.left_margin - section_g.right_margin
        agregar_tabulacion_derecha(p_header_g, ancho_g_emu)
    
        run_g_lbl = p_header_g.add_run("GRÁFICO")
        run_g_lbl.font.name = FUENTE
        run_g_lbl.font.size = Pt(TAMANO_BASE)
        run_g_lbl.font.bold = True
    
        p_header_g.add_run("\t")  # Salto de tabulador explícito para OXML
    
        run_pp_g = p_header_g.add_run("pp.")
        run_pp_g.font.name = FUENTE
        run_pp_g.font.size = Pt(TAMANO_BASE)
        run_pp_g.font.bold = True
    
        # Lista de Gráficos: data-driven desde contenido.py (GRAFICOS con campos "lista" y "pagina")
        graficos_def = [
            ("1", "Representación cartográfica y ubicación espacial de la empresa", "3"),
            ("2", "Organigrama estructural y niveles jerárquicos de la organización", "4"),
        ]
        graficos_cfg = getattr(c, 'GRAFICOS', [])
        if graficos_cfg:
            for g in graficos_cfg:
                num = str(g.get("numero", ""))
                desc = g.get("lista", g.get("titulo", "").split('. ', 1)[-1] if '. ' in g.get("titulo", "") else "")
                pag = str(g.get("pagina", ""))
                bookmark_id = f"bm_grafico{num}" if num else None
                agregar_fila_lista_preliminar_nativa(doc, num, desc, pag, bookmark_id=bookmark_id)
        else:
            for num, desc, pag in graficos_def:
                bookmark_id = f"bm_grafico{num}" if num else None
                agregar_fila_lista_preliminar_nativa(doc, num, desc, pag, bookmark_id=bookmark_id)
    
        # --- LISTA DE ANEXOS ---
        if hasattr(c, 'ANEXOS_LISTA') and c.ANEXOS_LISTA is not None:
            iniciar_seccion_preliminar(doc, "LISTA DE ANEXOS", bookmark_id="bm_lista_anexos")
            p_header_a = doc.add_paragraph()
            p_header_a.paragraph_format.space_before = Pt(12)
            p_header_a.paragraph_format.space_after = Pt(12)
        
            section_a = doc.sections[-1]
            ancho_a_emu = section_a.page_width - section_a.left_margin - section_a.right_margin
            agregar_tabulacion_derecha(p_header_a, ancho_a_emu)
        
            run_a_lbl = p_header_a.add_run("ANEXOS")
            run_a_lbl.font.name = 'Times New Roman'
            run_a_lbl.font.size = Pt(12)
            run_a_lbl.font.bold = True
        
            p_header_a.add_run("\t")  # Salto de tabulador explícito para OXML
        
            run_pp_a = p_header_a.add_run("pp.")
            run_pp_a.font.name = 'Times New Roman'
            run_pp_a.font.size = Pt(12)
            run_pp_a.font.bold = True
        
            anexos_lista = getattr(c, 'ANEXOS_LISTA', [])
            for idx, anexo in enumerate(anexos_lista):
                cod, desc = anexo[:2]
                letra = cod.split(" ")[-1]
                pag_est = str(13 + idx)
                bookmark_id = f"bm_anexo{letra.upper()}" if letra else None
                agregar_fila_lista_preliminar_nativa(doc, letra, desc, pag_est, bookmark_id=bookmark_id)
    
        # --- RESUMEN ---
        if hasattr(c, 'RESUMEN_TEXTO') and c.RESUMEN_TEXTO:
            iniciar_seccion_resumen(doc, c, bookmark_id="bm_resumen")
            p_res_texto = doc.add_paragraph()
            p_res_texto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_res_texto.paragraph_format.line_spacing = 1.15
            p_res_texto.paragraph_format.first_line_indent = Cm(1.25)
            p_res_texto.paragraph_format.space_after = Pt(6)
            run_res_texto = p_res_texto.add_run(c.RESUMEN_TEXTO)
            run_res_texto.font.name = FUENTE
            run_res_texto.font.size = Pt(TAMANO_BASE)
            p_kw = doc.add_paragraph()
            p_kw.paragraph_format.line_spacing = 1.15
            p_kw.paragraph_format.first_line_indent = Cm(1.25)
            run_kw_label = p_kw.add_run("Palabras claves: ")
            run_kw_label.font.bold = True
            run_kw_label.font.name = FUENTE
            run_kw_label.font.size = Pt(TAMANO_BASE)
            run_kw_texto = p_kw.add_run(c.PALABRAS_CLAVE)
            run_kw_texto.font.name = FUENTE
            run_kw_texto.font.size = Pt(TAMANO_BASE)

        # --- REGISTRO DEL INICIO DEL CUERPO ---
        idx_intro = len(doc.sections)
        iniciar_seccion_preliminar(doc, "INTRODUCCIÓN", bookmark_id="bm_introduccion", ocultar_primera=True)
        introduccion = getattr(c, 'INTRODUCCION_TEXTO', 'Texto de introducción no proporcionado.')
        if isinstance(introduccion, (list, tuple)):
            for parrafo in introduccion:
                agregar_parrafo_normado(doc, parrafo)
        else:
            agregar_parrafo_normado(doc, introduccion)

        # Retornamos el índice de la sección que se va a crear para el Capítulo I (que es la actual longitud de doc.sections)
        idx_cap1 = len(doc.sections)
    else:
        idx_cap1 = len(doc.sections)
        idx_intro = idx_cap1

# --- CAPÍTULO I: REALIDAD ORGANIZACIONAL ---
    iniciar_capitulo(doc, "I", "REALIDAD ORGANIZACIONAL", bookmark_id="bm_cap1")
    p_ident = doc.add_paragraph()
    p_ident.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ident.paragraph_format.space_before = ESP_DOBLE
    p_ident.paragraph_format.space_after = ESP_DOBLE
    p_ident.paragraph_format.line_spacing = INTERLINEADO
    r_ident = p_ident.add_run("Identificación de la empresa")
    r_ident.font.name = FUENTE
    r_ident.font.size = Pt(TAMANO_BASE)
    r_ident.font.bold = True
    _agregar_bookmark(p_ident, "bm_cap1_ident")

    agregar_titulo_nivel2(doc, "Razón social")
    agregar_parrafo_normado(doc, getattr(c, 'RAZON_SOCIAL', 'Razón Social no proporcionada.'), sangria=False)

    agregar_titulo_nivel2(doc, "Reseña histórica", bookmark_id="bm_cap1_resena")
    resena_data = getattr(c, 'RESENA_HISTORICA', [])
    if isinstance(resena_data, str):
        agregar_parrafo_normado(doc, resena_data)
    else:
        for parrafo in resena_data:
            agregar_parrafo_normado(doc, parrafo)

    agregar_titulo_nivel2(doc, "Misión", bookmark_id="bm_cap1_mision")
    agregar_parrafo_normado(doc, getattr(c, 'MISION', 'Misión no proporcionada.'))

    agregar_titulo_nivel2(doc, "Visión", bookmark_id="bm_cap1_vision")
    agregar_parrafo_normado(doc, getattr(c, 'VISION', 'Visión no proporcionada.'))

    agregar_titulo_nivel2(doc, "Valores", bookmark_id="bm_cap1_valores")
    agregar_parrafo_normado(doc, "Los valores que orientan las actividades de la organización destacan:")
    valores_data = getattr(c, 'VALORES', [])
    for i, (valor, descripcion) in enumerate(valores_data, 1):
        agregar_item_lista(doc, i, descripcion, valor)

    agregar_titulo_nivel2(doc, "Objetivos Organizacionales", bookmark_id="bm_cap1_obj")
    agregar_titulo_nivel2(doc, "Objetivo General")
    agregar_parrafo_normado(doc, getattr(c, 'OBJETIVO_GENERAL_EMPRESA', ''))
    objs_espec_emp = getattr(c, 'OBJETIVOS_ESPECIFICOS_EMPRESA', [])
    if objs_espec_emp:
        agregar_titulo_nivel2(doc, "Objetivos Específicos")
        for i, obj in enumerate(objs_espec_emp, 1):
            agregar_item_lista(doc, i, obj)

    carpeta_imagenes = getattr(c, 'CARPETA_IMAGENES', 'imagenes')
    _insertar_logo_empresa(doc, carpeta_imagenes)

    agregar_titulo_nivel2(doc, "Ubicación geográfica", bookmark_id="bm_cap1_ubic")
    agregar_parrafo_normado(doc, getattr(c, 'UBICACION', 'Ubicación no proporcionada.'), sangria=False)

    _insertar_graficos_por_ancla(doc, carpeta_imagenes, "ubicacion")

    agregar_titulo_nivel2(doc, "Estructura Organizativa", bookmark_id="bm_cap1_estruct")
    org_texto = getattr(c, 'ORGANIGRAMA_TEXTO', 'Estructura organizativa.')
    if isinstance(org_texto, list):
        for parrafo in org_texto:
            agregar_parrafo_normado(doc, parrafo)
    else:
        agregar_parrafo_normado(doc, org_texto)

    _insertar_graficos_por_ancla(doc, carpeta_imagenes, "estructura")

    # --- CAPÍTULO II: DIAGNÓSTICO SITUACIONAL ---
    if tiene_cap2:
        iniciar_capitulo(doc, "II", "DIAGNÓSTICO SITUACIONAL", bookmark_id="bm_cap2")
        agregar_titulo_nivel2(doc, "Identificación de la Situación Problemática", bookmark_id="bm_cap2_sit")
        situacion_problematica = getattr(c, 'SITUACION_PROBLEMATICA', [])
        if isinstance(situacion_problematica, str):
            agregar_parrafo_normado(doc, situacion_problematica)
        else:
            for bloque in situacion_problematica:
                if isinstance(bloque, dict):
                    titulo_bloque = bloque.get('titulo')
                    if titulo_bloque:
                        agregar_titulo_nivel2(doc, titulo_bloque)
                    parrafos_bloque = bloque.get('parrafos', [])
                    if isinstance(parrafos_bloque, str):
                        parrafos_bloque = [parrafos_bloque]
                    for parrafo in parrafos_bloque:
                        agregar_parrafo_normado(doc, parrafo)
                else:
                    agregar_parrafo_normado(doc, bloque)

        interrogante = getattr(c, 'INTERROGANTE_PROBLEMA', '')
        if interrogante:
            agregar_titulo_nivel2(doc, getattr(c, 'INTERROGANTE_TITULO', 'Interrogante orientadora'))
            agregar_parrafo_normado(doc, interrogante)

        agregar_titulo_nivel2(doc, "Objetivo General", bookmark_id="bm_cap2_objg")
        agregar_parrafo_normado(doc, getattr(c, 'OBJETIVO_GENERAL', 'Objetivo general no proporcionado.'))

        agregar_titulo_nivel2(doc, "Objetivos Específicos", bookmark_id="bm_cap2_obje")
        objs_especificos = getattr(c, 'OBJETIVOS_ESPECIFICOS', [])
        for i, obj in enumerate(objs_especificos, 1):
            agregar_item_lista(doc, i, obj)

        agregar_titulo_nivel2(doc, "Planificación integral de objetivos", bookmark_id="bm_cap2_planif")
        planif_datos = getattr(c, 'PLANIFICACION_DATOS', [])
        if planif_datos:
            p_planif_ant = doc.add_paragraph()
            intro_planif = getattr(c, 'PLANIFICACION_INTRO_TEXTO',
                "La planificación establece la relación entre cada objetivo y las actividades administrativas a ejecutar:")
            p_planif_ant = doc.add_paragraph()
            run_planif = p_planif_ant.add_run(intro_planif)
            run_planif.font.name = FUENTE
            run_planif.font.size = Pt(TAMANO_BASE)
            p_planif_ant.paragraph_format.line_spacing = INTERLINEADO
            titulo_planif = getattr(c, 'CUADRO_PLANIFICACION_TITULO', CUADRO_PLANIFICACION_TITULO_DEF)
            agregar_tabla_planificacion(doc, planif_datos, titulo_cuadro=titulo_planif, bookmark_id="bm_cuadro1")
            doc.add_paragraph()

        agregar_titulo_nivel2(doc, "Cronograma de actividades", bookmark_id="bm_cap2_crono")
        crono_datos = getattr(c, 'CRONOGRAMA_DATOS', [])
        if crono_datos:
            p_crono_ant = doc.add_paragraph()
            run_crono = p_crono_ant.add_run(getattr(c, 'CRONOGRAMA_INTRO_TEXTO',
                "El cronograma estructura temporalmente las tareas administrativas garantizando el cumplimiento del manual documental propuesto:"))
            run_crono.font.name = FUENTE
            run_crono.font.size = Pt(TAMANO_BASE)
            p_crono_ant.paragraph_format.line_spacing = INTERLINEADO

            # Sección horizontal (LANDSCAPE) de una sola hoja para el cuadro Gantt
            sec_land = doc.add_section(WD_SECTION_START.NEW_PAGE)
            sec_land.orientation = WD_ORIENT.LANDSCAPE
            sec_land.page_width = PAG_ALTO
            sec_land.page_height = PAG_ANCHO
            _config_seccion(sec_land, primera_sin_numero=False)
            # Márgenes más generosos en horizontal para que la tabla quepa en una sola hoja
            sec_land.left_margin = Cm(3)
            sec_land.right_margin = Cm(3)

            titulo_crono = getattr(c, 'CUADRO_CRONOGRAMA_TITULO', CUADRO_CRONOGRAMA_TITULO_DEF)
            agregar_gantt(doc, crono_datos, titulo_cuadro=titulo_crono, bookmark_id="bm_cuadro2")

            fuente_crono = getattr(c, 'CRONOGRAMA_FUENTE', '')
            if fuente_crono:
                p_fuente = doc.add_paragraph()
                p_fuente.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_fuente.paragraph_format.line_spacing = 1.0
                p_fuente.paragraph_format.space_before = Pt(6)
                p_fuente.paragraph_format.space_after = Pt(0)
                run_f = p_fuente.add_run(fuente_crono)
                run_f.font.name = FUENTE
                run_f.font.size = Pt(10)

            # Volver a vertical (PORTRAIT)
            sec_ret = doc.add_section(WD_SECTION_START.NEW_PAGE)
            sec_ret.orientation = WD_ORIENT.PORTRAIT
            sec_ret.page_width = PAG_ANCHO
            sec_ret.page_height = PAG_ALTO
            sec_ret.top_margin = MARGEN_SUP
            sec_ret.bottom_margin = MARGEN_INF
            sec_ret.left_margin = MARGEN_IZQ
            sec_ret.right_margin = MARGEN_DER
            _config_seccion(sec_ret, primera_sin_numero=False)
            doc.add_paragraph()

    # --- CAPÍTULO III: MARCO TEÓRICO ---
    if tiene_cap3:
        iniciar_capitulo(doc, "III", "MARCO TEÓRICO", bookmark_id="bm_cap3")
        if hasattr(c, 'BASES_TEORICAS') and isinstance(c.BASES_TEORICAS, list) and c.BASES_TEORICAS and isinstance(c.BASES_TEORICAS[0], dict):
            agregar_titulo_nivel2(doc, "Bases Teóricas Referenciales", bookmark_id="bm_cap3_bases")
            intro_mt = getattr(c, 'BASES_TEORICAS_INTRO', '')
            if intro_mt:
                agregar_parrafo_normado(doc, intro_mt)
            primer_sub = True
            for sub in c.BASES_TEORICAS:
                bm = "bm_cap3_bases" if primer_sub else None
                nivel = sub.get('nivel', 2)
                if nivel == 3:
                    agregar_titulo_nivel3(doc, sub.get('titulo', ''), bookmark_id=bm)
                elif nivel == 4:
                    agregar_titulo_nivel4(doc, sub.get('titulo', ''))
                else:
                    agregar_titulo_nivel2(doc, sub.get('titulo', ''), bookmark_id=bm)
                primer_sub = False
                for p in sub.get('parrafos', []):
                    agregar_parrafo_normado(doc, p)
                cita = sub.get('cita_larga')
                if cita and cita.get('texto'):
                    agregar_cita_larga(doc, cita['texto'], cita.get('autor', ''))
                    post_cita = sub.get('post_cita', getattr(c, 'POST_CITA_TEXTO', POST_CITA_TEXTO_DEF))
                    if post_cita:
                        agregar_parrafo_normado(doc, post_cita, sangria=True)
                posicion_autor = sub.get('posicion_autor')
                if posicion_autor:
                    agregar_parrafo_normado(doc, posicion_autor)
        else:
            agregar_titulo_nivel2(doc, "Bases Teóricas Referenciales", bookmark_id="bm_cap3_bases")
            bases_teoricas = getattr(c, 'BASES_TEORICAS_PARRAFOS', ['Bases teóricas referenciales.'])
            for parrafo in bases_teoricas:
                agregar_parrafo_normado(doc, parrafo)
            if hasattr(c, 'CITA_LARGA_TEXTO') and c.CITA_LARGA_TEXTO:
                agregar_cita_larga(doc, c.CITA_LARGA_TEXTO, getattr(c, 'CITA_LARGA_AUTOR', ''))
                post_cita = getattr(c, 'POST_CITA_TEXTO', POST_CITA_TEXTO_DEF)
                agregar_parrafo_normado(doc, post_cita, sangria=True)

    # --- CAPÍTULO IV: ACTIVIDADES REALIZADAS ---
    if tiene_cap4:
        iniciar_capitulo(doc, "IV", "ACTIVIDADES REALIZADAS", bookmark_id="bm_cap4")
        agregar_titulo_nivel2(doc, "Descripción de Actividades Ejecutadas por Semana", bookmark_id="bm_cap4_desc")
        agregar_parrafo_normado(doc, getattr(c, 'ACTIVIDADES_DESCRIPCION', 'Descripción de actividades ejecutadas.'))
        actividades_lista = getattr(c, 'ACTIVIDADES_LISTA', [])
        for i, actividad in enumerate(actividades_lista, 1):
            if isinstance(actividad, dict):
                agregar_titulo_nivel2(doc, f"Semana {actividad.get('semana', i)}", bookmark_id=f"bm_cap4_s{i}")
                if actividad.get('operativa'):
                    agregar_parrafo_seccion(doc, "Operativa", actividad['operativa'])
                if actividad.get('operativa_parrafo'):
                    agregar_parrafo_normado(doc, actividad['operativa_parrafo'])
                if actividad.get('investigacion'):
                    agregar_parrafo_seccion(doc, "De investigación", actividad['investigacion'])
                if actividad.get('investigacion_parrafo'):
                    agregar_parrafo_normado(doc, actividad['investigacion_parrafo'])
            else:
                agregar_item_lista(doc, i, actividad)

    # --- CAPÍTULO V: CONCLUSIONES Y RECOMENDACIONES ---
    if tiene_cap5:
        iniciar_capitulo(doc, "V", "CONCLUSIONES Y RECOMENDACIONES", bookmark_id="bm_cap5")
        agregar_titulo_nivel2(doc, "Conclusiones", bookmark_id="bm_cap5_concl")
        conclusiones_intro = getattr(c, 'CONCLUSIONES_INTRO', '')
        if conclusiones_intro:
            agregar_parrafo_normado(doc, conclusiones_intro)
        conclusiones = getattr(c, 'CONCLUSIONES', [])
        for i, conclusion in enumerate(conclusiones, 1):
            agregar_item_lista(doc, i, conclusion)

        agregar_titulo_nivel2(doc, "Recomendaciones", bookmark_id="bm_cap5_recom")
        recomendaciones = getattr(c, 'RECOMENDACIONES', [])
        if isinstance(recomendaciones, dict):
            rec_intro = recomendaciones.get('intro', '')
            if rec_intro:
                agregar_parrafo_normado(doc, rec_intro)
            for seccion in recomendaciones.get('secciones', []):
                agregar_titulo_nivel3(doc, seccion.get('titulo', ''))
                for item in seccion.get('items', []):
                    agregar_viñeta(doc, item)
        else:
            for i, recomendacion in enumerate(recomendaciones, 1):
                agregar_item_lista(doc, i, recomendacion)

    # --- REFERENCIAS BIBLIOGRÁFICAS ---
    iniciar_seccion_preliminar(doc, "REFERENCIAS", bookmark_id="bm_referencias")
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(24)
    p_sep.paragraph_format.space_after = Pt(0)
    referencias_lista = getattr(c, 'REFERENCIAS_LISTA', [])
    for ref in referencias_lista:
        agregar_referencia(doc, ref)

    # --- ANEXOS ---
    if hasattr(c, 'ANEXOS_LISTA') and c.ANEXOS_LISTA is not None:
        # Portadilla de ANEXOS (Art. 26: una hoja sola con la palabra ANEXOS centrada y en negrita)
        sec_portadilla = doc.add_section(WD_SECTION_START.NEW_PAGE)
        _config_seccion(sec_portadilla)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Empujar hacia el centro vertical aproximado de la página
        p.paragraph_format.space_before = ESP_PORTADILLA_ANEXOS
        run_anexos_tit = p.add_run("ANEXOS")
        run_anexos_tit.font.name = FUENTE
        _agregar_bookmark(p, "bm_anexos")
        run_anexos_tit.font.size = Pt(TAMANO_PORTADILLA_ANEXOS)
        run_anexos_tit.font.bold = True

    # Anexos individuales (Art. 15: cada uno en página nueva, arriba y centrado, subtítulo entre corchetes)
        for anexo in c.ANEXOS_LISTA:
            cod, desc = anexo[:2]
            numero_imagen = anexo[2] if len(anexo) > 2 else None
            alto_imagen = anexo[3] if len(anexo) > 3 else None
            contenido_anexo = anexo[4] if len(anexo) > 4 else None
            ancho_img = None
            alto_img = None
            if isinstance(alto_imagen, dict):
                ancho_img = alto_imagen.get('width_cm')
                alto_img = alto_imagen.get('height_cm')
            elif alto_imagen is not None:
                alto_img = alto_imagen
            sec_anexo = doc.add_section(WD_SECTION_START.NEW_PAGE)
            _config_seccion(sec_anexo)

            p_anexo = doc.add_paragraph()
            p_anexo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_anexo.paragraph_format.space_before = Pt(0)
            p_anexo.paragraph_format.space_after = ESP_DOBLE
            letra_anexo = cod.split(" ")[-1].upper()
            _agregar_bookmark(p_anexo, f"bm_anexo{letra_anexo}")
            
            # Nombre del Anexo (ej: ANEXO A)
            run_cod = p_anexo.add_run(cod.upper())
            run_cod.font.name = FUENTE
            run_cod.font.size = Pt(TAMANO_BASE)
            run_cod.font.bold = True
            
            # Subtítulo del contenido centrado entre corchetes [ ]
            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sub.paragraph_format.space_before = ESP_SENCILLO
            p_sub.paragraph_format.space_after = ESP_DOBLE
            
            run_desc = p_sub.add_run(f"[{desc}]")
            run_desc.font.name = FUENTE
            run_desc.font.size = Pt(TAMANO_BASE)
            run_desc.font.bold = True
            
            if numero_imagen is not None:
                ruta_imagen = buscar_imagen_por_numero(carpeta_imagenes, numero_imagen)
                if ruta_imagen:
                    p_imagen = doc.add_paragraph()
                    p_imagen.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if ancho_img is not None and alto_img is not None:
                        p_imagen.add_run().add_picture(ruta_imagen, width=Cm(ancho_img), height=Cm(alto_img))
                    elif alto_img is not None:
                        p_imagen.add_run().add_picture(ruta_imagen, height=Cm(alto_img))
                    elif ancho_img is not None:
                        p_imagen.add_run().add_picture(ruta_imagen, width=Cm(ancho_img))
                    else:
                        p_imagen.add_run().add_picture(ruta_imagen, width=Cm(14))

            if contenido_anexo:
                bloques = contenido_anexo if isinstance(contenido_anexo, (list, tuple)) else [contenido_anexo]
                for bloque in bloques:
                    agregar_parrafo_normado(doc, str(bloque))

    return idx_cap1, idx_intro, idx_indice_inicio, idx_indice_fin

# ================================================================
#  EJECUCIÓN PRINCIPAL
# ================================================================

def generar_reporte_completo(modo="completo"):
    print("» Inicializando documento...")
    doc = setup_iutecp_document()
    
    # 1. Portada (Sección 0) — solo autor, sin tutores
    construir_portada(doc, solo_autor=True, idx_seccion=0)
    
    # 2. Contraportada (Sección 1)
    sec_contra = doc.add_section(WD_SECTION_START.NEW_PAGE)
    sec_contra.top_margin = Cm(3)
    sec_contra.bottom_margin = Cm(3)
    sec_contra.left_margin = Cm(4)
    sec_contra.right_margin = Cm(3)
    sec_contra.different_first_page_header_footer = True
    construir_portada(doc, idx_seccion=1)
    
    # 3. Cuerpo (Sección 2 o 1 en borrador)
    idx_cap1, idx_intro, idx_indice_inicio, idx_indice_fin = construir_cuerpo_documento(doc, modo=modo)
    
    # Aplicar la numeración de página correcta en base al índice dinámico del Capítulo I
    agregar_numeracion_pie(doc, idx_inicio_cuerpo=idx_intro, idx_indice_inicio=idx_indice_inicio, idx_indice_fin=idx_indice_fin)

    if modo == "borrador1":
        sufijo = "_BORRADOR1"
    elif modo == "borrador2":
        sufijo = "_BORRADOR2"
    elif modo == "borrador3":
        sufijo = "_BORRADOR3"
    elif modo == "borrador4":
        sufijo = "_BORRADOR4"
    else:
        sufijo = ""
    docx_output = f"Informe_Pasantia_IUTECP{sufijo}.docx"
    doc.save(docx_output)
    print(f"✔ Archivo Word generado: {docx_output}")

    print("» Renderizando PDF usando LibreOffice...")
    pdf_output = f"Informe_Pasantia_IUTECP{sufijo}.pdf"
    soffice_cmd = 'libreoffice'
    if platform.system() == 'Windows':
        possible_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        for path in possible_paths:
            if os.path.exists(path):
                soffice_cmd = path
                break

    try:
        subprocess.run([soffice_cmd, '--headless', '--convert-to', 'pdf', docx_output], check=True, stdout=subprocess.DEVNULL)
        print("✔ ¡PDF generado con éxito!")
    except FileNotFoundError:
        print("❌ Error: LibreOffice no está instalado o no se encontró en el PATH.")
    except subprocess.CalledProcessError as e:
        print(f"❌ Error en la conversión a PDF: {e}")

if __name__ == "__main__":
    import sys
    modo = "completo"
    if "--modo" in sys.argv:
        idx = sys.argv.index("--modo")
        if idx + 1 < len(sys.argv):
            modo = sys.argv[idx + 1]
    generar_reporte_completo(modo=modo)
